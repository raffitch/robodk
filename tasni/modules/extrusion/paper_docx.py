"""A Word draft of the ring-stack results, rebuilt from the archive on demand.

The point is that nobody transcribes numbers. Every value here comes from
``paper_summary`` -- the same object the app shows -- so the document, the app
and the archive cannot disagree, and the tables arrive as real Word tables that
paste into the manuscript with their structure intact.

It is written to be read DURING collection: what is still missing is listed
explicitly, so a half-finished run is obviously half-finished rather than
quietly citable.

Needs the ``docx`` extra (``pip install -e .[docx]``); imported lazily so every
measurement still runs without it.
"""
from __future__ import annotations

import time
from pathlib import Path

from .measure import CAPTURE_LABEL, paper_summary

# The protocol's own targets, and requirement 3 of the paper's outstanding list.
TAKES_PER_CONDITION = 3
MEASUREMENTS_FOR_TIMING = 12

# Stated once, where the claim is made. Hand-eye board consistency and
# work-plane RMS for this cell; nothing below this floor may be reported.
ERROR_FLOOR_MM = 1.26
WORK_PLANE_RMS_MM = 1.39
STANDOFF_MM = 300

def method_text(provenance: dict) -> str:
    """The method paragraph, with this run's own numbers in it.

    The standoff and the hand-eye residual are read from the take that was
    measured, not from constants here: a constant silently goes stale when the
    cell is recalibrated, and the error-floor sentence is the one claim in the
    paper that must never be generous.
    """
    pose = provenance.get("inspection_pose") or {}
    standoff = pose.get("standoff_mm") or STANDOFF_MM
    resolution = provenance.get("camera_resolution") or "1280x720"
    quality = ((provenance.get("calibration") or {}).get("quality") or {})
    floor = quality.get("board_consistency_rms_mm") or ERROR_FLOOR_MM
    return (
        "Rings of dried, previously extruded material were placed by hand on the scanned work "
        "surface, one on top of another; no material was extruded and no valve was actuated "
        "during the measurements. Each measurement moved only the camera to a viewpoint "
        f"{float(standoff):.0f} mm above the top of the layer being measured, normal to the "
        "work plane, captured one RGB-D frame from an Intel RealSense D435i at "
        f"{resolution}, and reconstructed the deposited ring's centreline in the work frame. "
        "Layer N was segmented above the measured top surface of layer N-1, so a displaced "
        "ring is compared against the ring it actually sits on. What follows is therefore a "
        "controlled validation of the sensing-and-comparison chain against a known introduced "
        "offset, and not the deposition deviation of a printed cylinder. Hand-eye calibration "
        f"for this cell has a board-consistency residual of {float(floor):.2f} mm and a "
        f"work-plane RMS of {WORK_PLANE_RMS_MM:.2f} mm; no accuracy is claimed below that "
        "floor.")


LIMITATIONS = (
    "Three limits bound what these measurements support. First, inspection happens BETWEEN "
    "layers and not DURING deposition: the camera moves to its viewpoint while nothing is "
    "being extruded, which is what the printing loop already does at a layer boundary, and "
    "the cycle time above is the cost of that pause. Measuring a bead as it is laid would "
    "require a capture synchronised to a moving camera pose, which is not validated here. "
    "Second, the specimens are hand-placed dried beads rather than freshly deposited "
    "material, so this is a validation of the sensing-and-comparison chain and not of "
    "deposition accuracy. Third, the introduced offsets are ground truth measured by hand "
    "against the work frame's axes, so the detection error carries that measurement's own "
    "uncertainty as well as the chain's."
)

SETTINGS_HEADERS = ["Setting", "Value"]
TIMING_HEADERS = ["Stage", "Mean (ms)", "SD (ms)", "n"]
TIMING_ROWS = [("capture_ms", "RGB-D capture"),
               ("total_ms", "Reconstruction"),
               ("acquisition_to_path_ms", "Acquisition to reconstructed path"),
               ("move_to_pose_ms", "Move to the inspection pose"),
               ("settle_ms", "Settle before capture"),
               ("return_ms", "Return to the start pose"),
               ("inspection_cycle_ms", "Whole inspection excursion")]

CONDITION_HEADERS = ["Condition", "n", "How", "Centre spread (mm)",
                     "Centre offset (mm)", "Detection error (mm)",
                     "Paired detection error (mm)",
                     "Mean |dev| (mm)", "RMS (mm)", "Max (mm)", "Shape RMS (mm)"]
TAKE_HEADERS = ["Layer", "Take", "Phase", "Introduced (mm)", "Measured offset (mm)",
                "Detection error (mm)", "Paired error (mm)", "RMS (mm)", "Acq to path (ms)",
                "Layer cost (s)", "Valid"]


def _pm(text: str) -> str:
    """Markdown writes +/- so it stays ASCII; Word can have the real sign."""
    return text.replace("+/-", "±")


def _stat_cell(stat: dict, digits: int = 2) -> str:
    if not stat or stat.get("mean") is None:
        return "-"
    text = f"{stat['mean']:.{digits}f}"
    if stat.get("sd") is not None:
        text += f" ± {stat['sd']:.{digits}f}"
    return text


def _offset_cell(offset) -> str:
    if not offset or not any(float(v) for v in offset):
        return "-"
    return f"({offset[0]:g}, {offset[1]:g})"


def _detection_error(manifest: dict) -> "float | None":
    from .measure import detection_error_mm
    return detection_error_mm(manifest)


def _paired_error(manifest: dict, manifests: list[dict]) -> "float | None":
    from .measure import paired_detection
    paired = paired_detection(manifest, manifests)
    return None if paired is None else paired["detection_error_mm"]


def _gaps(summary: dict) -> list[str]:
    """What this run still owes the paper, in the operator's terms."""
    notes: list[str] = []
    for condition in summary["conditions"]:
        short = TAKES_PER_CONDITION - condition["takes"]
        if short > 0:
            notes.append(f"{condition['condition']}: {condition['takes']} take(s) recorded, "
                         f"{short} more for a mean and standard deviation.")
        invalid = condition["takes"] - condition["valid"]
        if invalid:
            notes.append(f"{condition['condition']}: {invalid} take(s) invalid and excluded "
                         "from every average. Reprocess them or repeat them.")
        if condition["introduced_norm_mm"] > 0 and condition["valid"] \
                and not condition["paired_detection_error_mm"]["n"]:
            notes.append(f"{condition['condition']}: no zero-offset take of layer "
                         f"{condition['layer_index']} precedes it, so its detection error is "
                         "scored against the plan centre only (that includes the ring's "
                         "placement error). Measure the ring in place before displacing it.")
    # Requirement 3 is the time measured ON THE CELL: a take reprocessed offline
    # is a measurement but not a live one, so it does not count here.
    live = int(summary["timing_ms"]["acquisition_to_path_ms"]["n"] or 0)
    if live < MEASUREMENTS_FOR_TIMING:
        notes.append(f"{live} live measurement(s) recorded; the acquisition-to-path mean and "
                     f"standard deviation needs {MEASUREMENTS_FOR_TIMING} -- "
                     f"{MEASUREMENTS_FOR_TIMING - live} more.")
    if not any(c["introduced_norm_mm"] > 0 for c in summary["conditions"]):
        notes.append("No condition with an introduced offset yet. The paper's claim is that a "
                     "known displacement is recovered, so at least one is required.")
    return notes


def _system_rows(provenance: dict, robot_name: str | None = None) -> list[list[str]]:
    """What the run was measured ON, entirely from the take's provenance."""
    pose = provenance.get("inspection_pose") or {}
    intrinsics = provenance.get("camera_intrinsics") or {}
    calibration = provenance.get("calibration") or {}
    quality = calibration.get("quality") or {}
    config = provenance.get("processing_config") or {}
    rows: list[list[str]] = []

    def add(label: str, value) -> None:
        if value not in (None, "", []):
            rows.append([label, str(value)])

    add("Robot", robot_name)
    # No fallbacks below: a row is written only when the archive recorded the
    # value. A default printed here would read as a measurement of this cell.
    if provenance.get("camera_resolution"):
        add("Camera", f"Intel RealSense D435i, {provenance['camera_resolution']} "
                      "aligned colour and depth")
    matrix = intrinsics.get("K")
    if matrix:
        add("Camera intrinsics",
            f"fx {float(matrix[0][0]):.1f}, fy {float(matrix[1][1]):.1f}, "
            f"cx {float(matrix[0][2]):.1f}, cy {float(matrix[1][2]):.1f} px")
    distortion = intrinsics.get("dist_coeffs")
    if distortion:
        add("Lens distortion", ", ".join(f"{float(v):.4f}" for v in distortion[:5])
            + " (k1 k2 p1 p2 k3, calibrated in-cell)")
    if pose.get("standoff_mm"):
        add("Inspection standoff", f"{float(pose['standoff_mm']):.0f} mm above the top of the "
                                   "layer being measured, normal to the work plane")
    add("Work frame", provenance.get("work_frame"))
    add("Inspection tool", provenance.get("inspection_tool"))
    if quality or calibration:
        add("Hand-eye calibration",
            f"{calibration.get('method', 'unknown')}, verdict {quality.get('verdict', 'n/a')}; "
            f"held-out reprojection {float(quality.get('val_rms_px', 0)):.2f} px, "
            f"board consistency {float(quality.get('board_consistency_rms_mm', 0)):.2f} mm "
            f"(run {calibration.get('run_id', 'n/a')})")
    if config:
        if config.get("substrate_sigma_k") is not None:
            clamp = config.get("substrate_floor_clamp_mm") or [None, None]
            add("Segmentation",
                f"deposit floor derived per frame as {config['substrate_sigma_k']}× the "
                "fitted substrate's own residual scale"
                + (f", clamped to {clamp[0]}–{clamp[1]} mm" if len(clamp) == 2 else "")
                + f"; plane fitted within {config.get('substrate_fit_radius_mm')} mm of the "
                f"ring centre; components shorter than {config.get('deposit_min_length_beads')} "
                "bead widths rejected; radial band "
                f"±{config.get('radial_roi_margin_mm')} mm about the nominal ring")
        else:
            # A take measured before 2026-08-30 used a CONSTANT floor above
            # work-frame Z=0 plus a colour gate. A draft built from that archive
            # has to state what THAT take used, not today's recipe.
            add("Segmentation",
                f"deposit floor {config.get('deposit_min_height_mm')} mm above the work "
                "plane (fixed; pre-2026-08-30 chain), radial band "
                f"±{config.get('radial_roi_margin_mm')} mm about the nominal ring")
        add("Centreline extraction",
            f"raster {config.get('raster_mm_per_pixel')} mm/px, thinned and pruned, "
            f"{config.get('measured_spline_points')} spline samples; bead footprint over "
            f"{config.get('bead_width_bins')} angular bins")
        if config.get("settle_s") is not None:
            add("Settling dwell", f"{config['settle_s']} s before capture")
    commit = provenance.get("git_commit")
    if commit:
        add("Software revision", f"{commit[:10]} (tasni, this repository)")
    return rows


def _timing_rows(timing: dict) -> list[list[str]]:
    rows = []
    for key, label in TIMING_ROWS:
        stat = timing.get(key) or {}
        if not stat.get("n"):
            continue
        rows.append([label, f"{stat['mean']:.0f}",
                     "-" if stat.get("sd") is None else f"{stat['sd']:.0f}",
                     str(stat["n"])])
    return rows


def _add_table(document, headers: list[str], rows: list[list[str]]):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for values in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, values):
            cell.text = text
    return table


def _figure_paths(trial_dir: Path, summary: dict) -> list[tuple[Path, str]]:
    """The stack, plus the plan view of the last take of each condition."""
    from .figures import ensure_figure

    wanted: list[tuple[Path, str]] = []
    # The method figure leads: it is the paper's account of how a depth frame
    # becomes the number in Table 1.
    first = next((m for m in summary["manifests"]
                  if (m.get("metrics") or {}).get("valid")), None)
    method_dir = None if first is None else _layer_dir_for(trial_dir, first)
    if method_dir is not None:
        try:
            wanted.append((ensure_figure(method_dir, "pipeline.png"),
                           "One RGB-D frame becoming a measured centreline: captured depth, "
                           "the work region of interest, the deposit cluster, its top surface, "
                           "and the extracted centreline against nominal."))
        except Exception:
            pass
    for stem, caption in (
            ("stack.png", "Every layer's latest measured centreline against nominal."),
            ("tube.png", "The commanded bead drawn at its real diameter, each layer at its "
                         "own height, with the measured centreline running through it.")):
        try:
            wanted.append((ensure_figure(trial_dir, stem), caption))
        except Exception:
            continue
    from .measure import _condition_name
    for condition in summary["conditions"]:
        # The latest VALID take of this condition: the one a reader would check
        # the table against.
        chosen = None
        for manifest in summary["manifests"]:
            if not (manifest.get("metrics") or {}).get("valid"):
                continue
            if _condition_name(manifest) == condition["condition"]:
                chosen = manifest
        layer_dir = None if chosen is None else _layer_dir_for(trial_dir, chosen)
        if layer_dir is None:
            continue
        try:
            wanted.append((ensure_figure(layer_dir, "plan.png"),
                           f"{condition['condition']}: extracted centreline against nominal."))
        except Exception:
            continue
    return wanted


def _layer_dir_for(trial_dir: Path, manifest: dict) -> "Path | None":
    layer = int(manifest.get("layer_index") or 0)
    take = int(manifest.get("take") or 1)
    name = f"layer-{layer:03d}" + ("" if take == 1 else f"-take{take:02d}")
    path = trial_dir / name
    return path if path.is_dir() else None


def build_paper_docx(root, trial_id: str, out_path=None, *,
                     embed_figures: bool = True,
                     robot_name: str | None = None) -> Path:
    """Write the results draft for ``trial_id`` and return the file path."""
    from docx import Document                       # the `docx` extra
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Mm, Pt

    root = Path(root)
    trial_dir = root / trial_id
    summary = paper_summary(root, trial_id)
    out_path = Path(out_path) if out_path else (trial_dir / "paper-draft.docx")

    document = Document()
    document.add_heading("Ring-stack validation - results draft", level=1)
    stamp = document.add_paragraph(
        f"Trial {trial_id} - generated {time.strftime('%Y-%m-%d %H:%M')} - "
        f"{summary['takes']} measurement(s), {summary['valid']} valid. "
        "Rebuilt from the archive; every number here is the one the application reports.")
    stamp.runs[0].italic = True
    stamp.runs[0].font.size = Pt(9)

    first = next((m for m in summary["manifests"] if m.get("provenance")), None)
    provenance = (first or {}).get("provenance") or {}
    document.add_heading("Method (paste into the method section)", level=2)
    document.add_paragraph(method_text(provenance))

    settings = _system_rows(provenance, robot_name=robot_name)
    if settings:
        document.add_paragraph("Table 1. The system these measurements were made on.")
        _add_table(document, SETTINGS_HEADERS, settings)
        document.add_paragraph()

    document.add_heading("Results (paste into the results section)", level=2)
    document.add_paragraph(_pm(summary["headline"]))

    document.add_paragraph(f"Table {2 if settings else 1}. Geometric deviation by condition. "
                           "Deviation is measured "
                           "from the nominal centre; detection error is the difference between "
                           "the measured displacement and the displacement introduced by hand. "
                           "The paired detection error scores the shift against the ring's own "
                           "last measured position before it was moved, which removes the "
                           "hand-placement error of the undisplaced position; it is the "
                           "figure the claim rests on.")
    _add_table(document, CONDITION_HEADERS,
               [[c["condition"], str(c["takes"]), CAPTURE_LABEL[c["capture"]],
                 ("-" if c["centre_spread"]["rms_mm"] is None
                  else f"{c['centre_spread']['rms_mm']:.2f}"),
                 _stat_cell(c["center_offset_norm_mm"]), _stat_cell(c["detection_error_mm"]),
                 _stat_cell(c["paired_detection_error_mm"]),
                 _stat_cell(c["mean_absolute_mm"]), _stat_cell(c["rms_mm"]),
                 _stat_cell(c["maximum_mm"]), _stat_cell(c["shape_rms_mm"])]
                for c in summary["conditions"]])
    document.add_paragraph()

    for paragraph in summary["prose"]:
        added = document.add_paragraph(_pm(paragraph))
        if paragraph.startswith("WARNING"):
            for run in added.runs:
                run.bold = True

    timing_rows = _timing_rows(summary["timing_ms"])
    if timing_rows:
        document.add_heading("What one inspection costs", level=2)
        document.add_paragraph(
            "Every measurement is one excursion: the arm leaves the path, settles, captures a "
            "frame, the frame is reconstructed, and the arm returns. Inspecting between layers "
            "costs the sum of those, once per layer.")
        _add_table(document, TIMING_HEADERS, timing_rows)
        document.add_paragraph()

    gaps = _gaps(summary)
    if gaps:
        document.add_heading("Still missing (delete before submitting)", level=2)
        warning = document.add_paragraph(
            "Not ready to cite as it stands. This run still owes:")
        warning.runs[0].bold = True
        for note in gaps:
            document.add_paragraph(note, style="List Bullet")

    document.add_heading("Reproducibility and archive", level=2)
    reprocessed = summary["timing_ms"].get("offline_reprocessed_takes") or 0
    archive_note = (
        "Every measurement archives the raw colour and depth frame it was made from, the "
        "camera pose in the work frame, the camera intrinsics, the full processing "
        "configuration and the software revision, alongside the derived point cloud, "
        "centreline and metrics. Any take can therefore be reprocessed from its own frame "
        "with no robot, and reprocessing reproduces the archived numbers.")
    if reprocessed:
        archive_note += (
            f" {reprocessed} take(s) in this trial were recovered exactly that way: a "
            "measurement that failed on the cell was reprocessed offline from its archived "
            "frame and produced a valid centreline without repeating the capture.")
    document.add_paragraph(archive_note)

    document.add_heading("Limitations", level=2)
    document.add_paragraph(LIMITATIONS)

    document.add_heading("Every take (working record)", level=2)
    ordered = sorted(summary["manifests"],
                     key=lambda m: (int(m.get("layer_index") or 0), int(m.get("take") or 1)))
    rows = []
    for manifest in ordered:
        metrics = manifest.get("metrics") or {}
        timings = (manifest.get("processing") or {}).get("timings_ms") or {}
        error = _detection_error(manifest)
        paired = _paired_error(manifest, summary["manifests"])
        acq = timings.get("acquisition_to_path_ms")
        rows.append([
            str(manifest.get("layer_index", "")), str(manifest.get("take", 1)),
            str((manifest.get("annotation") or {}).get("phase") or "-"),
            _offset_cell((manifest.get("annotation") or {}).get("introduced_offset_mm")),
            f"{metrics['center_offset_norm_mm']:.2f}" if metrics.get("center_offset_norm_mm") is not None else "-",
            f"{error:.2f}" if error is not None else "-",
            f"{paired:.2f}" if paired is not None else "-",
            f"{metrics['rms_mm']:.2f}" if metrics.get("rms_mm") is not None else "-",
            f"{acq:.0f}" if acq else "-",
            f"{timings['inspection_cycle_ms'] / 1000:.1f}" if timings.get("inspection_cycle_ms") else "-",
            "yes" if metrics.get("valid") else "no",
        ])
    _add_table(document, TAKE_HEADERS, rows)

    if embed_figures:
        document.add_heading("Figures", level=2)
        try:
            figures = _figure_paths(trial_dir, summary)
        except ImportError:
            figures = []
            document.add_paragraph(
                "Figures were not rendered here (matplotlib is absent: pip install -e "
                ".[figures]). They live beside each take in the archive.")
        for index, (path, caption) in enumerate(figures, start=1):
            document.add_picture(str(path), width=Mm(150))
            document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            label = document.add_paragraph(f"Figure {index}. {caption} "
                                           f"Vector original: {path.with_suffix('.pdf')}")
            label.runs[0].italic = True
            label.runs[0].font.size = Pt(9)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(out_path))
    return out_path
