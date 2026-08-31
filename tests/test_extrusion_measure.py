"""Ring-stack measure-only experiment: synthetic proof, processing, jobs, API."""
from __future__ import annotations

import json
import math
import re
import sys
from pathlib import Path

import numpy as np
import pytest

sys.path.insert(0, str(Path(__file__).resolve().parent))

import extrusion_synthetic as syn  # noqa: E402
import geometry_fixtures as gf  # noqa: E402
from tasni.modules.extrusion.processing import depth_to_work_points  # noqa: E402


def test_renderer_puts_a_ring_where_it_says_at_the_height_it_says():
    center = (200.0, 150.0)
    T = syn.inspection_camera_T([center[0], center[1], 6.0], 300.0)
    depth = syn.render_scene([syn.RingSpec(60.0, 8.0, center, height_fn=syn.flat(6.0))], T,
                             plane_center_xy_mm=center, noise_mm=0.0)
    assert depth.dtype == np.uint16 and depth.shape == (720, 1280)
    points, raw = depth_to_work_points(depth, syn.geometry(), T)
    assert raw > 100_000                                  # plane + ring both rendered
    ring = points[points[:, 2] > 3.0]
    radii = np.linalg.norm(ring[:, :2] - np.array(center), axis=1)
    assert 55.0 < radii.min() and radii.max() < 65.0     # 60 +/- bead/2 (+ rounding)
    assert 5.0 < ring[:, 2].max() < 7.0                   # crest at 6 mm
    plane = points[points[:, 2] <= 1.0]
    assert len(plane) > 50_000


# ---------------------------------------------------------------- circle metrics

from tasni.modules.extrusion.comparison import compare_circle


def test_shifted_circle_reports_its_offset_and_zero_shape_error():
    theta = np.linspace(0, 2 * np.pi, 360, endpoint=False)
    shifted = np.column_stack((40 * np.cos(theta) + 10, 40 * np.sin(theta), np.full(360, 5.0)))
    m = compare_circle(shifted, 40.0, nominal_center_mm=(0.0, 0.0))
    assert m.center_offset_mm == pytest.approx((10.0, 0.0), abs=1e-6)
    assert m.center_offset_norm_mm == pytest.approx(10.0, abs=1e-6)
    assert m.shape_rms_mm < 1e-6 and m.shape_max_mm < 1e-6
    # Deviation is still measured from the NOMINAL centre (the paper's number).
    assert m.mean_absolute_mm == pytest.approx(6.35, abs=0.05)
    assert m.rms_mm == pytest.approx(7.06, abs=0.05)
    assert m.maximum_mm == pytest.approx(10.0, abs=0.05)


def test_old_metrics_payload_without_offset_fields_still_validates():
    from tasni.modules.extrusion.models import DeviationMetrics
    old = DeviationMetrics(mean_absolute_mm=1, rms_mm=1, maximum_mm=1,
                           measured_center_mm=(0, 0), measured_radius_mm=40,
                           path_completeness=1, maximum_angular_gap_deg=2, valid=True)
    assert old.center_offset_norm_mm == 0.0 and old.shape_rms_mm == 0.0


# ------------------------------------------------- end-to-end synthetic processing

from tasni.core.config import ExtrusionConfig
from tasni.modules.extrusion.inspection import aim_point_mm
from tasni.modules.extrusion.models import CylinderRecipe, CylinderSetup
from tasni.modules.extrusion.processing import process_observation
from tasni.modules.extrusion.toolpath import generate_cylinder_plan, points_array

CENTER = (200.0, 150.0)


def scene_plan(*, radius=60.0, bead=8.0, layers=1, layer_height=6.0, center=CENTER):
    recipe = CylinderRecipe(radius_mm=radius, layer_count=layers, layer_height_mm=layer_height,
                            bead_diameter_mm=bead, robot_speed_mm_s=75, extrusion_rate_pct=0,
                            points_per_circle=180)
    setup = CylinderSetup(print_tool="LongCalibTool", work_frame="Tasni Work Frame",
                          inspection_tool="Realsense", inspection_auto=True,
                          center_x_mm=center[0], center_y_mm=center[1])
    return generate_cylinder_plan(recipe, setup)


def observe(plan, layer_index, rings, *, config=None, seed=0, stages=None):
    """Render the rings from the derived inspection pose and process that frame."""
    layer = plan.layers[layer_index - 1]
    T = syn.inspection_camera_T(aim_point_mm(plan.recipe, plan.setup, layer_index), 300.0)
    depth = syn.render_scene(rings, T, plane_center_xy_mm=(plan.setup.center_x_mm,
                                                           plan.setup.center_y_mm), seed=seed)
    kwargs = {}
    if stages is not None:
        kwargs["stages"] = stages
    return process_observation(depth=depth,
                               geometry=syn.geometry(),
                               T_work_camera=T,
                               plan=plan, layer=layer, config=config or ExtrusionConfig(),
                               **kwargs)


def test_true_ring_measures_as_true():
    pytest.importorskip("open3d")
    plan = scene_plan()
    out = observe(plan, 1, [syn.RingSpec(60.0, 8.0, CENTER, height_fn=syn.flat(6.0))])
    m = out.metrics
    assert m.valid, m.warnings
    assert m.mean_absolute_mm < 1.0 and m.rms_mm < 1.0
    assert abs(m.measured_radius_mm - 60.0) < 1.0
    assert m.center_offset_norm_mm < 1.0
    assert m.path_completeness >= 0.95
    assert out.report["timings_ms"]["total_ms"] > 0


def test_ring_shifted_10mm_reports_the_shift():
    pytest.importorskip("open3d")
    plan = scene_plan()
    out = observe(plan, 1, [syn.RingSpec(60.0, 8.0, (CENTER[0] + 10.0, CENTER[1]),
                                          height_fn=syn.flat(6.0))])
    m = out.metrics
    assert m.valid, m.warnings
    assert m.center_offset_mm[0] == pytest.approx(10.0, abs=1.0)
    assert abs(m.center_offset_mm[1]) < 1.0
    assert m.maximum_mm == pytest.approx(10.0, abs=1.5)
    assert m.mean_absolute_mm == pytest.approx(6.36, abs=1.0)
    assert m.rms_mm == pytest.approx(7.06, abs=1.0)
    assert m.shape_rms_mm < 1.0


def _board_bias_patch(center, *, r_from: float, r_to: float, half_height_mm: float,
                      z_mm: float, step_mm: float = 0.25) -> np.ndarray:
    """A patch of the build plane reading a few mm HIGH, touching the ring's outer flank.

    What the D435i does to the ChArUco board at 300 mm: broad patches biased by
    2-5 mm (measured 2026-08-28: bare board z p50 0.8 / p99 4.8 mm, 22.7% above the
    2.5 mm deposit floor). Flat, so its normals face straight up, and fused to the
    ring, so it lands in the ring's DBSCAN cluster.

    ``step_mm`` is the SOURCE-sample spacing, and it defaults to 0.25 mm to match
    :meth:`extrusion_synthetic.RingSpec.surface_points` -- not a detail, a
    faithfulness requirement. ``render_depth`` is a point splatter with no
    surface interpolation: it can only fill the depth pixels a source sample
    actually lands on, so a source grid coarser than the depth pixel pitch
    renders a SIEVE. At these scenes' 300 mm standoff one colour pixel spans
    304.5 / 889.87 = 0.342 mm on the patch, so the 1.0 mm spacing this defaulted
    to until 2026-08-31 filled only 12.2% of the patch's own footprint in the
    depth image (measured: 1131 source points -> 1131 of 9296 pixels; 0.5 mm ->
    47.2%; 0.25 mm -> 9296 of 9296, 100.0%). Real board depth bias is solid, and
    a holey patch is a materially EASIER adversary -- it voxel-downsamples to
    fewer points, biases the trim's circle fit less, and dilates into a thinner
    lobe. That is not a cosmetic difference: measured on the guard in
    :func:`test_board_depth_bias_fused_to_the_ring_does_not_break_the_measurement`
    (limit r 72.0 mm), the holey 1.0 mm patch kept the raster inside 70.6 mm and
    passed, while the same scene rendered solid leaked board out to 72.4 mm
    (0.5 mm sampling) and 72.8-73.4 mm (0.25 mm) -- i.e. the fixture was passing
    because of its own sampling artefact, not because the chain rejected the
    contamination.

    So: do NOT raise this back toward (or past) the depth-pixel pitch. A fixture
    that renders contamination the sensor could not produce cannot test the
    chain's ability to reject the contamination it does produce.
    """
    xs = np.arange(center[0] + r_from, center[0] + r_to + step_mm, step_mm)
    ys = np.arange(center[1] - half_height_mm, center[1] + half_height_mm + step_mm, step_mm)
    X, Y = np.meshgrid(xs, ys, indexing="ij")
    return np.column_stack((X.ravel(), Y.ravel(), np.full(X.size, float(z_mm))))


def test_the_board_bias_fixture_renders_a_solid_patch_not_a_sieve():
    """The guard on the guard, and the reason it is a test and not a comment.

    Every board-contamination assertion below is only worth what the fixture's
    realism is worth, and that realism is invisible: ``_board_bias_patch``
    returns a perfectly regular grid of points at any spacing, and the hole it
    leaves at a coarse spacing appears only after ``render_depth`` splats it.
    Until 2026-08-31 the default was 1.0 mm and the rendered patch was 12%
    full -- and the r 72.0 mm guard below passed BECAUSE of that, not despite
    it. So pin the property that matters directly: at the default sampling the
    patch fills every depth pixel of its own footprint, the way real board depth
    bias does. The second half pins the trap itself, so a future edit that
    coarsens the spacing fails here, naming the reason, instead of quietly
    turning the adversary back into a sieve.
    """
    T = syn.inspection_camera_T([CENTER[0], CENTER[1], 6.0], 300.0)

    def fill_fraction(**kwargs):
        patch = _board_bias_patch(CENTER, r_from=62.0, r_to=100.0,
                                  half_height_mm=14.0, z_mm=3.5, **kwargs)
        depth = syn.render_depth(patch, T, noise_mm=0.0)
        rows, cols = np.nonzero(depth)
        footprint = ((cols.max() - cols.min() + 1) * (rows.max() - rows.min() + 1))
        return int(np.count_nonzero(depth)) / footprint

    assert fill_fraction() >= 0.999, (
        "the default board patch renders holey -- a sieve is an easier adversary "
        "than the solid depth bias the real board produces (it downsamples to "
        "fewer points, biases the trim's circle fit less, and dilates into a "
        "thinner lobe), so every rejection assertion built on it is worth less "
        "than it looks; see _board_bias_patch's docstring")
    # One colour pixel spans 0.342 mm on this patch, so 1.0 mm source spacing
    # cannot fill it -- measured 12.2%. This is the artefact, kept on record.
    assert fill_fraction(step_mm=1.0) < 0.2


def observe_with_board_bias(plan, layer_index, rings, patch, *, seed=0, config=None):
    layer = plan.layers[layer_index - 1]
    T = syn.inspection_camera_T(aim_point_mm(plan.recipe, plan.setup, layer_index), 300.0)
    centre = (plan.setup.center_x_mm, plan.setup.center_y_mm)
    parts = [syn.plane_points(center_xy_mm=centre), patch]
    parts += [ring.surface_points() for ring in rings]
    depth = syn.render_depth(np.vstack(parts), T, seed=seed)
    return process_observation(depth=depth,
                               geometry=syn.geometry(),
                               T_work_camera=T,
                               plan=plan, layer=layer,
                               config=config or ExtrusionConfig())


def test_board_depth_bias_fused_to_the_ring_does_not_break_the_measurement():
    """Cell 2026-08-28 20:48: a board patch at ~3.5 mm, touching the ring, dilated
    into a lobe with a 37 mm skeleton arm and exhausted the branch guard.

    The fix has to work by SHAPE -- the patch is inside the bead's own height band
    (bead z p25 1.8 / p50 3.8 mm), so no floor separates them without also
    destroying the ring (a 3 mm floor read r 36.7 for a 42.6 mm ring, and passed).
    """
    pytest.importorskip("open3d")
    plan = scene_plan()
    ring = syn.RingSpec(60.0, 8.0, CENTER, height_fn=syn.flat(6.0))
    patch = _board_bias_patch(CENTER, r_from=62.0, r_to=100.0, half_height_mm=14.0, z_mm=3.5)

    out = observe_with_board_bias(plan, 1, [ring], patch)

    m = out.metrics
    assert m.valid, m.warnings
    assert m.measured_radius_mm == pytest.approx(60.0, abs=1.0)
    assert m.center_offset_norm_mm < 1.0
    assert m.path_completeness >= 0.95
    r = np.linalg.norm(np.asarray(out.filtered_xyz)[:, :2] - np.asarray(CENTER), axis=1)
    assert r.max() < 60.0 + 8.0 + 4.0, "board points beyond the bead must not reach the raster"
    assert out.report["counts"]["after_radial_trim"] < out.report["counts"]["after_largest_cluster"]


@pytest.mark.parametrize("voxel_size_m", [0.001, 0.0005])
def test_a_solid_board_patch_is_kept_out_of_the_raster_at_either_voxel(voxel_size_m):
    """The guard above, held at BOTH voxel sizes the chain runs at.

    The test above pins one (config, sampling) pair. This one pins the property
    across the voxel size, because that default is not frozen -- it is 1.0 mm
    today and moving to 0.5 mm -- and the leak this file was written to catch
    was present at BOTH. Measured 2026-08-31 on the solid (0.25 mm) patch,
    before the ``_radial_trim`` convergence fix:

        voxel 1.0 mm -> r.max 72.821    voxel 0.5 mm -> r.max 73.378

    against the same 72.0 mm limit, i.e. the contamination reached the raster
    regardless of how finely the cloud was downsampled -- as it must, since the
    leak is the trim's own circle fit still moving when the band schedule runs
    out, and a voxel size cannot fix that. A fix that holds at one voxel and not
    the other is therefore not a fix; it is a coincidence, and this test is what
    stops one being mistaken for the other.
    """
    pytest.importorskip("open3d")
    plan = scene_plan()
    ring = syn.RingSpec(60.0, 8.0, CENTER, height_fn=syn.flat(6.0))
    patch = _board_bias_patch(CENTER, r_from=62.0, r_to=100.0, half_height_mm=14.0, z_mm=3.5)

    out = observe_with_board_bias(plan, 1, [ring], patch,
                                  config=ExtrusionConfig(voxel_size_m=voxel_size_m))

    m = out.metrics
    assert m.valid, m.warnings
    assert m.measured_radius_mm == pytest.approx(60.0, abs=1.0)
    assert m.center_offset_norm_mm < 1.0
    r = np.linalg.norm(np.asarray(out.filtered_xyz)[:, :2] - np.asarray(CENTER), axis=1)
    assert r.max() < 60.0 + 8.0 + 4.0, (
        f"voxel {voxel_size_m * 1000:g} mm: board points beyond the bead reached "
        f"the raster (r.max {r.max():.3f} mm)")


def test_radial_trim_follows_a_displaced_ring_not_the_nominal():
    """The trim is about the FITTED circle: a ring shifted 12 mm must still read 12."""
    pytest.importorskip("open3d")
    plan = scene_plan()
    moved = (CENTER[0] + 12.0, CENTER[1])
    ring = syn.RingSpec(60.0, 8.0, moved, height_fn=syn.flat(6.0))
    patch = _board_bias_patch(moved, r_from=62.0, r_to=100.0, half_height_mm=14.0, z_mm=3.5)

    out = observe_with_board_bias(plan, 1, [ring], patch)

    m = out.metrics
    assert m.valid, m.warnings
    assert m.center_offset_mm[0] == pytest.approx(12.0, abs=1.0)
    assert m.measured_radius_mm == pytest.approx(60.0, abs=1.0)


def test_real_frame_with_board_noise_measures_the_applied_ring():
    """The 2026-08-28 20:48 cell frame that failed with `branch guard exhausted`."""
    pytest.importorskip("open3d")
    fixture = np.load(Path(__file__).parent / "fixtures" / "extrusion" / "ring2"
                      / "ring2_board_noise_20260828.npz")
    centre = tuple(float(v) for v in fixture["nominal_center_mm"])
    plan = scene_plan(radius=float(fixture["recipe_radius_mm"]),
                      bead=float(fixture["recipe_bead_mm"]),
                      layer_height=float(fixture["recipe_layer_height_mm"]),
                      center=centre)
    depth = fixture["depth"]

    out = process_observation(
        depth=depth,
        geometry=gf.aligned(fixture["K"], (depth.shape[1], depth.shape[0])),
        T_work_camera=fixture["T_work_camera"],
        plan=plan, layer=plan.layers[0], config=ExtrusionConfig())

    m = out.metrics
    assert m.valid, m.warnings
    # Characterize, minutes earlier on the same ring: r 42.60, centre (214.64, 146.69).
    assert m.measured_radius_mm == pytest.approx(42.6, abs=1.0)
    assert m.measured_center_mm == pytest.approx(centre, abs=2.0)
    assert m.path_completeness >= 0.95
    assert m.shape_rms_mm < 2.5
    r = np.linalg.norm(np.asarray(out.filtered_xyz)[:, :2] - np.asarray(centre), axis=1)
    assert not (r > 55.0).any(), "the board lobe at r 55-72 mm must be gone (was 21% of points)"


# ------------------------- the board patch dies in GEOMETRY, where colour used to

RING1_TAKE04 = (Path(__file__).parent / "fixtures" / "extrusion" / "ring1"
                / "ring1_take04_branchguard_20260829.npz")


def _ring1_take04() -> dict:
    """Cell trial 20260829-165938, layer 1 take 4: the frame that crashed.

    The archive still carries this take's colour frame (``color_jpeg``); the
    chain no longer reads it, and neither do these tests. It stays as evidence:
    it is the frame the 20:1 saturation separation was measured on, and the one
    that shows the separation had inverted a day later.
    """
    fixture = np.load(RING1_TAKE04)
    centre = tuple(float(v) for v in fixture["nominal_center_mm"])
    return {
        "depth": fixture["depth"], "K": fixture["K"],
        "T_work_camera": fixture["T_work_camera"], "centre": centre,
        "plan": scene_plan(radius=float(fixture["recipe_radius_mm"]),
                           bead=float(fixture["recipe_bead_mm"]),
                           layer_height=float(fixture["recipe_layer_height_mm"]),
                           center=centre),
    }


def _measure_ring1_take04(**overrides):
    f = _ring1_take04()
    # This fixture is a pre-protocol-2 capture with 1 mm depth WORDS, so it is
    # processed at the 2 mm voxel it was captured under. The 1 mm default sits
    # at this archive's quantisation floor, where it merges nothing and lets
    # noise through -- it flips this frame's branch-guard outcome in BOTH
    # directions across the two archived takes. On protocol-2 depth (0.1 mm
    # words) 1 mm spans ten quantisation steps, which is the point. Shared by
    # both halves of the pair below: their whole claim, "the patch dies in
    # geometry", only holds if both measure the SAME frame at the SAME voxel.
    return f, process_observation(
        depth=f["depth"], geometry=gf.aligned(f["K"], (1280, 720)),
        T_work_camera=f["T_work_camera"], plan=f["plan"], layer=f["plan"].layers[0],
        config=ExtrusionConfig(voxel_size_m=0.002, **overrides))


def test_the_board_patch_dies_by_geometry_not_colour():
    """The 2026-08-29 13:03 cell frame, through the chain that reads no colour.

    A 22-point patch of bare black checker, 12 mm outside the ring's +X flank,
    welded to the bead. The colour gate used to remove it; nothing reads colour
    now, so the removal has to come from shape and from a floor derived from the
    frame's own substrate noise. MEASURED 2026-08-31 running this fixture through
    the new chain (task 7 step 5a) -- every number below is that measurement, not
    a target:

      * The frame is MEASURED, not refused, and reported **INVALID**: path
        completeness 0.846, maximum angular gap 55.3 deg. That is the honest
        outcome. This is a 1 mm-quantised capture, so the substrate's own sigma
        is 0.759 mm and 3*sigma saturates the clamp at 2.0 mm -- which lands at
        ~2.54 mm in work-frame Z, reproducing what the old 2.5 mm floor did to
        this ring (fixture README: completeness 0.87, 46 deg gap). The
        low-relief sector whose crest reads 2.9-4.9 mm falls out, and the
        metrics say so instead of returning a ring that was not there.
        Protocol-2 captures (0.1 mm words, sigma ~0.55) derive a 1.65 mm floor
        and keep that sector -- see tests/test_extrusion_golden.py.
      * Of the 1478 patch points reaching the work ROI, 7 survive to the crest,
        and the radius bias is down from the +0.6-0.7 mm takes 1-3 carried to
        +0.24 mm. **That cleaning is NOT compactness's doing**, and the two
        facts must not be welded into one claim. Run with the filter ON and OFF
        (measured 2026-08-31):

            stage            ON              OFF
            work_roi         n=22226 p=1478  n=22226 p=1478
            compactness      n=21714 p=1386  n=22226 p=1478
            deposit_cluster  n= 1266 p= 105  n= 1277 p= 116
            radial_trimmed   n= 1231 p=  72  n= 1228 p=  72
            top_surface      n=  432 p=   7  n=  436 p=   6

        Compactness removes 92 of the 1471 patch points that go away -- 6% --
        and the pre-existing downstream chain (DBSCAN, the radial trim about the
        fitted circle, the crest filter) reaches the same endpoint either way:
        identically 72 after the trim, 6 crest points without the filter against
        7 with it.
      * What compactness actually does here is drop five compact components
        totalling 512 points, which changes the RASTER TOPOLOGY -- and that is
        what decides the branch-guard outcome. Without it this frame reproduces
        the 2026-08-29 cell abort, which is what the twin below pins. That role
        is real and load-bearing; it is simply a different mechanism from
        "compactness cleans the patch off the crest". The filter's
        contamination-rejection value is separately evidenced on the archive
        (spec §3.5: layer-001 keeps 1 of 1 component, layer-002 keeps 2 of 4 and
        1 of 9, the rejected ones being exactly the compact patches).

    The one outcome the design could not accept -- a VALID measurement that
    still includes the patch, i.e. a silently wrong radius -- is not what this
    frame produces, and ``assert not m.valid`` is what pins that.
    """
    pytest.importorskip("open3d")
    f, out = _measure_ring1_take04()

    m = out.metrics
    # Honest refusal-by-metric: no wrong number is presented as a good one.
    assert not m.valid
    assert m.path_completeness == pytest.approx(0.846, abs=0.03)
    assert m.maximum_angular_gap_deg > 30.0
    assert any("completeness" in w or "angular gap" in w for w in m.warnings), m.warnings

    sub = out.report["substrate"]
    assert sub["source"] == "fitted_plane"
    assert sub["sigma_mm"] == pytest.approx(0.759, abs=0.05)
    # 3 * 0.759 = 2.28 -> clamped. The clamp is what stops a noisy frame opening
    # the floor to everything; here it also costs the low-relief sector.
    assert sub["floor_mm"] == 2.0
    # ... and the fitted plane is the BOARD, 1.3 mm below work Z=0 (spec §1).
    assert sub["plane"][2] == pytest.approx(-1.32, abs=0.2)

    # Five of six components dropped, and the fail-open bypass did NOT fire
    # (a real rejection, not a starvation). This is what changes the raster
    # topology; the crest cleaning is the downstream chain's -- see the
    # ON/OFF table in the docstring.
    assert sub["compactness"]["compactness_components"] == 6
    assert sub["compactness"]["compactness_kept_components"] == 1
    assert sub["compactness"]["compactness_bypassed"] == 0

    cluster = np.asarray(out.filtered_xyz)
    r = np.linalg.norm(cluster[:, :2] - np.asarray(f["centre"]), axis=1)
    # The patch sat on the ring's +X flank at r 50-54 mm (raw, out to r 71 mm).
    patch = (cluster[:, 0] > 254.0) & (r > 47.0)
    assert int(patch.sum()) <= 12, "the +X board patch must be all but gone"
    assert r.max() < 54.0, "nothing may survive past the bead's own outer flank"
    # Characterization of this ring, minutes earlier: r 42.2 mm.
    assert m.measured_radius_mm == pytest.approx(42.2, abs=1.0)


def test_the_same_frame_reproduces_the_cell_crash_with_compactness_disabled():
    """Locks the CAUSE, exactly as the colour-gate pair used to.

    With ``deposit_min_length_beads = 0`` the five compact components the filter
    would have dropped (512 points) stay in the cloud, and the raster topology
    they produce exhausts the branch guard -- the original cell abort,
    reproduced (measured 2026-08-31: branch pixels 1, 1, 2 across the three
    attempts against a 15 px spur limit).

    Note what this does and does not show. It is NOT that those points reach the
    crest and are measured as bead: the crest is cleaned to 6-7 patch points
    either way (see the ON/OFF table in the test above). It is that their
    presence changes which skeleton the thinning produces, and the guard refuses
    the branched one. That is the load-bearing role compactness inherited from
    the saturation gate, and it stays measured rather than asserted.

    Guards against 'fixing' this by loosening the branch guard instead: the
    guard was right, and takes 1-3 carried the same contamination into radii
    biased 0.6-0.7 mm large precisely because their topology let them through.
    """
    pytest.importorskip("open3d")
    with pytest.raises(RuntimeError, match="branch guard exhausted"):
        _measure_ring1_take04(deposit_min_length_beads=0.0)


# ------------------------------- branch-guard spur limit: measured, not nominal

def test_spur_guard_uses_the_deposits_own_footprint_not_an_inflated_recipe_bead():
    """A caller's ``recipe.bead_diameter_mm`` can be stale or simply wrong -- a
    fresh trial's generic default, an operator's guess before any ring has been
    characterized. The spur-pruning tolerance must not inherit that number
    uncritically: an OVERSTATED bead inflates the tolerance, which prunes a
    LONGER twig without asking whether it is real (see the next test for what
    that costs). Two calls on the SAME clean ring, one with the true bead and
    one with it more than doubled, must (a) both measure the ring the same way
    -- this is a topology guard, not a measurement stage -- and (b) both report
    a ``spur_guard_bead_mm`` clamped near the true footprint, not the inflated
    recipe value.
    """
    pytest.importorskip("open3d")
    T = syn.inspection_camera_T([CENTER[0], CENTER[1], 6.0], 300.0)
    ring = syn.RingSpec(40.0, 10.0, CENTER, height_fn=syn.flat(6.0))
    scene = np.vstack((syn.plane_points(center_xy_mm=CENTER), ring.surface_points()))
    depth = syn.render_depth(scene, T, noise_mm=0.3)

    def run(bead_mm):
        plan = scene_plan(radius=40.0, bead=bead_mm, layer_height=6.0, center=CENTER)
        return process_observation(
            depth=depth, geometry=syn.geometry(),
            T_work_camera=T, plan=plan, layer=plan.layers[0],
            config=ExtrusionConfig())

    true_bead = run(10.0)          # recipe already matches the physical ring
    inflated = run(25.0)           # recipe overstates the bead by 2.5x

    assert true_bead.report["valid"] and inflated.report["valid"]
    true_guard_mm = true_bead.report["counts"]["spur_guard_bead_mm"]
    inflated_guard_mm = inflated.report["counts"]["spur_guard_bead_mm"]
    assert true_guard_mm == pytest.approx(9.5, abs=1.5)
    # The whole point: NOT the inflated 25 mm the caller supplied.
    assert inflated_guard_mm < 15.0
    assert inflated_guard_mm == pytest.approx(true_guard_mm, abs=1.0)
    # The measurement itself is unmoved by the recipe's (wrong) bead guess.
    assert inflated.metrics.measured_radius_mm == pytest.approx(
        true_bead.metrics.measured_radius_mm, abs=1.0)
    assert (inflated.report["geometry"]["bead_width_mean_mm"]
            == pytest.approx(true_bead.report["geometry"]["bead_width_mean_mm"], abs=1.0))


def test_spur_guard_still_catches_real_contamination_regardless_of_recipe_bead():
    """The guard must not go silent just because a caller's recipe happens to be
    wrong in the SAFE direction either. A ring with a genuine tangential shelf
    of extra material welded to its outer edge (the synthetic stand-in for the
    2026-08-29 cell's board-patch-on-the-flank failure) has to keep tripping
    the branch guard whether the recipe's bead assumption is accurate or
    grossly inflated -- clamping the tolerance to the measured footprint must
    never let a genuinely contaminated topology through just because the
    caller supplied a bad number.
    """
    pytest.importorskip("open3d")
    T = syn.inspection_camera_T([CENTER[0], CENTER[1], 6.0], 300.0)
    ring = syn.RingSpec(40.0, 10.0, CENTER, height_fn=syn.flat(6.0))
    # A tangential shelf just proud of the ring's own outer edge, spanning a
    # 28 degree arc -- long enough that neither an accurate nor a doubled+
    # dilation/spur tolerance can absorb it into the clean loop.
    #
    # 28, not the original 20: once the synthetic fixtures moved to protocol-2
    # depth words the crest resolves cleanly enough that a 20 degree shelf
    # (16.8 mm of arc, against a 15-16 px spur limit) sits ON the guard's
    # threshold -- it refused at a 10 mm recipe bead and was absorbed at 30 mm,
    # differing by a single pixel of tolerance. Measured 2026-08-31 across
    # spans 20/25/28/30/32/35 deg at recipe beads 10/20/30 mm: 25-30 deg
    # refuses at EVERY recipe bead (a plateau, not an edge), 20 deg is
    # marginal, and 32 deg and beyond is a concentric arc the radial trim
    # removes cleanly (measured r 39.96-40.00 for a 40 mm ring). 28 deg sits
    # mid-plateau, so this pins the guard rather than a rounding.
    r0 = 40.0 + 10.0 / 2.0 + 3.0
    thetas = np.deg2rad(np.arange(0.0, 28.0, 0.5))
    radii = np.arange(r0 - 2.5, r0 + 2.5, 0.5)
    Th, R = np.meshgrid(thetas, radii, indexing="ij")
    shelf = np.column_stack((
        CENTER[0] + R.ravel() * np.cos(Th.ravel()),
        CENTER[1] + R.ravel() * np.sin(Th.ravel()),
        np.full(Th.size, 4.0)))
    scene = np.vstack((syn.plane_points(center_xy_mm=CENTER), ring.surface_points(), shelf))
    depth = syn.render_depth(scene, T, noise_mm=0.3)

    for bead_mm in (10.0, 20.0):
        plan = scene_plan(radius=40.0, bead=bead_mm, layer_height=6.0, center=CENTER)
        with pytest.raises(RuntimeError, match="branch guard exhausted"):
            process_observation(
                depth=depth, geometry=syn.geometry(),
                T_work_camera=T, plan=plan, layer=plan.layers[0],
                config=ExtrusionConfig())


@pytest.mark.xfail(strict=True, reason=(
    "KNOWN DEFECT, pinned deliberately: a 3x overstated recipe bead inflates "
    "_rasterize's DILATION (radius = bead/2), which fattens the ring and the "
    "shelf into one lobe whose thinned skeleton has no junction left, so the "
    "branch guard cannot see the contamination it is there to catch. Measured "
    "2026-08-31: r 39.04-39.14 for a 40.0 mm ring, VALID, across shelf spans "
    "20-32 deg -- a confident number ~0.9 mm wrong. The spur TOLERANCE is "
    "already clamped to the frame's own measured footprint (`a0fabca`) for "
    "exactly this reason; the raster's dilation is not. Feeding it the same "
    "clamped bead was tried and reverted: it fixes this (39.87) but regresses "
    "the real 2026-08-28 ring2 frame, where the tighter dilation stops "
    "absorbing the ChArUco board lobe and exhausts the guard on a take that "
    "measures correctly today. Closing this needs its own evidence on real "
    "frames. When it is closed this test XPASSes and strict=True fails the "
    "suite -- delete the marker then. Latent before the synthetic fixtures "
    "moved to protocol-2 depth words: a coarser, noisier crest happened to "
    "leave a branch behind, so the guard fired for the wrong reason."))
def test_an_overstated_recipe_bead_defeats_the_branch_guard_through_the_raster():
    pytest.importorskip("open3d")
    T = syn.inspection_camera_T([CENTER[0], CENTER[1], 6.0], 300.0)
    ring = syn.RingSpec(40.0, 10.0, CENTER, height_fn=syn.flat(6.0))
    r0 = 40.0 + 10.0 / 2.0 + 3.0
    thetas = np.deg2rad(np.arange(0.0, 20.0, 0.5))
    radii = np.arange(r0 - 2.5, r0 + 2.5, 0.5)
    Th, R = np.meshgrid(thetas, radii, indexing="ij")
    shelf = np.column_stack((
        CENTER[0] + R.ravel() * np.cos(Th.ravel()),
        CENTER[1] + R.ravel() * np.sin(Th.ravel()),
        np.full(Th.size, 4.0)))
    scene = np.vstack((syn.plane_points(center_xy_mm=CENTER), ring.surface_points(), shelf))
    depth = syn.render_depth(scene, T, noise_mm=0.3)
    plan = scene_plan(radius=40.0, bead=30.0, layer_height=6.0, center=CENTER)
    with pytest.raises(RuntimeError, match="branch guard exhausted"):
        process_observation(depth=depth, geometry=syn.geometry(), T_work_camera=T,
                            plan=plan, layer=plan.layers[0], config=ExtrusionConfig())


def test_branch_guard_exhaustion_names_the_tolerance_it_gave_up_with():
    """An abort has to say WHICH tolerance it gave up with, and where it came from.

    Since the spur tolerance was clamped to the frame's own measured bead
    (``a0fabca``) there are two very different reasons this can raise, with
    opposite remedies: a genuinely contaminated frame (leave the guard alone --
    it is doing its job) or a clean ring whose measured bead came in narrow
    enough to drop ``spur_limit`` a pixel below what this frame needed. The
    number that separates them is ``spur_bead_mm``, and it was computed and then
    thrown away: it is recorded in ``counts["spur_guard_bead_mm"]``, but
    ``counts`` only reaches a caller through the report and this path raises
    before a report exists. So on the one occasion the number decides what to do
    next -- a take that aborted at the cell -- nobody could see it.

    Guarded by a test because the message is the entire diagnostic: the raw
    RGB-D is archived and the take can be reprocessed, so what an abort costs is
    the operator's time working out which of the two cases they are in.
    """
    pytest.importorskip("open3d")
    T = syn.inspection_camera_T([CENTER[0], CENTER[1], 6.0], 300.0)
    ring = syn.RingSpec(40.0, 10.0, CENTER, height_fn=syn.flat(6.0))
    # 28 degrees: see test_spur_guard_still_catches_real_contamination... for
    # the span sweep that measured 25-30 deg as the plateau where this shelf
    # refuses at every recipe bead.
    r0 = 40.0 + 10.0 / 2.0 + 3.0
    thetas = np.deg2rad(np.arange(0.0, 28.0, 0.5))
    radii = np.arange(r0 - 2.5, r0 + 2.5, 0.5)
    Th, R = np.meshgrid(thetas, radii, indexing="ij")
    shelf = np.column_stack((
        CENTER[0] + R.ravel() * np.cos(Th.ravel()),
        CENTER[1] + R.ravel() * np.sin(Th.ravel()),
        np.full(Th.size, 4.0)))
    scene = np.vstack((syn.plane_points(center_xy_mm=CENTER), ring.surface_points(), shelf))
    depth = syn.render_depth(scene, T, noise_mm=0.3)

    # The recipe deliberately overstates the bead 2x, so the clamped value and the
    # recipe value are far apart and the message cannot pass by quoting one twice.
    # (2x, not 3x: at 3x the raster's dilation absorbs the contamination and the
    # guard never fires at all -- see
    # test_an_overstated_recipe_bead_defeats_the_branch_guard_through_the_raster.)
    plan = scene_plan(radius=40.0, bead=20.0, layer_height=6.0, center=CENTER)
    cfg = ExtrusionConfig()
    with pytest.raises(RuntimeError, match="branch guard exhausted") as excinfo:
        process_observation(
            depth=depth, geometry=syn.geometry(),
            T_work_camera=T, plan=plan, layer=plan.layers[0],
            config=cfg)

    message = str(excinfo.value)
    limit = re.search(r"spur_limit (\d+) px", message)
    assert limit, message
    clamped = re.search(r"1\.5 x ([\d.]+) mm bead", message)
    assert clamped, message

    # The tolerance quoted must be the one actually applied -- the same formula the
    # loop prunes with, on the CLAMPED bead, not the recipe's inflated 30 mm.
    spur_bead_mm = float(clamped.group(1))
    assert int(limit.group(1)) == max(2, int(math.ceil(
        1.5 * spur_bead_mm / cfg.raster_mm_per_pixel)))
    assert spur_bead_mm < 20.0, message           # clamped, not the recipe's 30 mm

    # And the two inputs that decide which remedy applies are both named, so the
    # reader can see the clamp fired rather than having to infer it.
    assert "recipe bead 20.000 mm" in message, message
    assert re.search(r"this frame measured [\d.]+ mm", message), message


def test_the_derived_floor_lands_where_the_constant_used_to():
    """clamp(k * sigma) must land in the old constant's neighbourhood on the
    synthetic plane (spec §3.4: 1.55-1.74 mm measured on the cell archive), and
    the report must carry the §4 health block."""
    pytest.importorskip("open3d")
    plan = scene_plan(radius=40.0, bead=9.0, layer_height=6.0)
    out = observe(plan, 1, [syn.RingSpec(40.0, 9.0, CENTER, height_fn=syn.flat(6.0))])
    sub = out.report["substrate"]
    assert sub["source"] == "fitted_plane"
    assert sub["sigma_mm"] > 0.0
    assert 1.0 <= sub["floor_mm"] <= 2.0
    assert out.metrics.valid


def test_the_substrate_block_reports_everything_a_frame_must_say_about_itself():
    """Spec §4: one health block per take, or a degrading setup shows up only as
    a low completeness that reads like a placement fault.

    ``compactness_kept_components`` vs ``compactness_components`` is load-bearing
    on its own: the filter's fail-open is a STARVATION guard, not a partial-ring
    guard, so a real ring fragmented into arcs can lose most of its fragments
    with ``compactness_bypassed = 0`` and NOTHING else in the report would say
    so (measured on layer-002-take03 of the 2026-08-30 archive: 6 components in,
    1 kept).
    """
    pytest.importorskip("open3d")
    plan = scene_plan(radius=40.0, bead=9.0, layer_height=6.0)
    out = observe(plan, 1, [syn.RingSpec(40.0, 9.0, CENTER, height_fn=syn.flat(6.0))])
    sub = out.report["substrate"]
    assert set(sub) >= {"source", "sigma_mm", "tilt_deg", "plane", "inlier_fraction",
                        "bias_correction_mm", "bias_correction_sigma", "floor_mm",
                        "plane_offset_at_center_mm", "substrate_p99_mm",
                        "separation_margin_mm", "compactness"}
    assert set(sub["compactness"]) == {"compactness_components",
                                       "compactness_kept_components",
                                       "compactness_bypassed"}
    assert sub["compactness"]["compactness_kept_components"] >= 1
    assert sub["compactness"]["compactness_bypassed"] == 0
    assert len(sub["plane"]) == 3
    # Separation margin = bead p50 - substrate p99. A 6 mm ring on a quiet
    # synthetic plane has metres of headroom; what is pinned is that it is
    # POSITIVE and derived, not that it hits a particular value.
    assert sub["separation_margin_mm"] > 1.0
    # substrate_p99_mm is NOT asserted below the floor: it is measured on the
    # substrate uncensored, precisely so it CAN report a tail that has climbed
    # into the deposit band (it does, on the archive's layer-2 takes). Asserting
    # it below the floor is what the old censored form guaranteed by
    # construction, which is why that assertion said nothing.
    assert sub["substrate_p99_mm"] > 0.0
    # The old chain's `report["floor"]` block is gone, not renamed alongside.
    assert "floor" not in out.report


def test_the_shape_gate_rejection_names_the_capture_that_produced_the_blobs():
    """A rejection has to distinguish a leaking front end from a genuinely low ring.

    2026-08-30 16:35, cell: candidate 1 was 5489 points at r 51.6 mm with a 79 mm
    radial span -- a disc, around a physical ring measured at r 42.3 mm. The
    message listed all eighteen blobs and their failed criteria and said nothing
    about where they came from, so the two opposite remedies (fix the capture /
    re-place the ring) were indistinguishable from it. Same defect the branch
    guard's message had until `da5f7a4`, same fix: this function RAISES, so the
    counts the capture already measured die with it unless they ride along.

    The counts that carry that diagnosis are the SEGMENTATION's own now -- the
    derived floor and the substrate sigma it came from -- where they used to be
    the colour gate's kept-fraction. Both answer the same question: did the front
    end admit far more of the frame than the ring?
    """
    from tasni.modules.extrusion.processing import _select_ring_cluster

    rng = np.random.default_rng(0)
    radius = 60.0 * np.sqrt(rng.random(4000))
    theta = rng.random(4000) * 2 * math.pi
    disc = np.column_stack((250.0 + radius * np.cos(theta),
                            140.0 + radius * np.sin(theta),
                            rng.random(4000) * 3.0))       # filled: span ratio ~1.5
    counts = {"substrate_floor_mm": 1.5, "substrate_sigma_mm": 0.5,
              "search_cylinder_above_floor_fraction": 0.1871}

    with pytest.raises(RuntimeError, match="shape gate") as excinfo:
        _select_ring_cluster([disc], np.array([212.1, 149.7]), counts)

    message = str(excinfo.value)
    payload = json.loads(message[message.index("{"):])
    assert payload["candidates"][0]["radial_span_ratio"] > 0.8      # the gate was right
    assert payload["capture"]["substrate_sigma_mm"] == pytest.approx(0.5)
    assert payload["capture"]["search_cylinder_above_floor_fraction"] == pytest.approx(0.1871)
    assert payload["capture"]["substrate_floor_mm"] == pytest.approx(1.5)


def test_characterize_records_what_the_search_cylinder_held_before_the_floor():
    """The floor is derived; the noise it has to beat is a fact of the frame.

    Characterization fits the same substrate the layer measurements do -- it
    defines the recipe they are judged against, so the two must see the same
    cloud. The floor it derives, and the fraction of the search cylinder that
    clears it, cannot be read from ``after_search_roi`` alone (on 2026-08-30
    18.7% of the cylinder cleared the old fixed 1.5 mm), so the frame states
    both.
    """
    pytest.importorskip("open3d")
    plan = scene_plan(radius=60.0, bead=8.0, layer_height=5.0)
    T = syn.inspection_camera_T(aim_point_mm(plan.recipe, plan.setup, 1), 300.0)
    depth = syn.render_scene([syn.RingSpec(60.0, 8.0, CENTER, height_fn=syn.flat(6.0))],
                             T, plane_center_xy_mm=CENTER)
    found = characterize_ring(depth=depth,
                              geometry=syn.geometry(),
                              T_work_camera=T,
                              search_center_mm=CENTER, work_frame="Tasni Work Frame",
                              config=ExtrusionConfig())

    coarse = found.report["counts_coarse"]
    # Derived from THIS frame's substrate noise, inside the configured clamp --
    # never the 1.5/2.5 mm constants the colour gate used to pick between.
    assert 1.0 <= coarse["substrate_floor_mm"] <= 2.0
    assert coarse["substrate_sigma_mm"] > 0.0
    assert coarse["search_cylinder_points"] > coarse["after_search_roi"]
    assert 0.0 < coarse["search_cylinder_above_floor_fraction"] < 1.0
    # The statistic describes the WHOLE cylinder, not the deposit the ROI kept:
    # its median is the substrate and only its tail is the ring's 6 mm crest.
    # That is what makes it readable as "the board is in here too". Heights, not
    # work-frame Z: the band above is applied to heights.
    assert coarse["search_cylinder_height_mm_p50"] == pytest.approx(0.0, abs=1.0)
    assert coarse["search_cylinder_height_mm_p99"] > 4.0


def test_the_previous_layer_floor_is_gone_by_design():
    """Spec §2.4: previous-layer referencing made the only stacked data WORSE
    (0.62 -> 0.50). Layers are measured against the substrate; the ROI ceiling
    accommodates the stack. Its return needs new evidence, not a revert."""
    import inspect
    from tasni.modules.extrusion.processing import process_observation
    assert "floor_profile" not in inspect.signature(process_observation).parameters


# ------------------------------------------------------ ring geometry (Task 5)

def test_wavy_ring_height_profile_is_measured():
    pytest.importorskip("open3d")
    plan = scene_plan(layer_height=7.5)
    out = observe(plan, 1, [syn.RingSpec(60.0, 8.0, CENTER, height_fn=syn.wavy(7.5, 2.5, lobes=2))])
    g = out.geometry
    # The datum is the surface FITTED IN THIS FRAME, never the work frame's Z=0.
    assert g is not None and g.height_reference == "fitted_plane"
    assert g.top_z_min_mm == pytest.approx(5.0, abs=1.5)
    assert g.top_z_max_mm == pytest.approx(10.0, abs=1.5)
    assert g.top_z_std_mm > 1.0
    assert g.height_mean_mm == pytest.approx(7.5, abs=1.0)


def test_bead_width_is_the_rings_radial_footprint():
    pytest.importorskip("open3d")
    plan = scene_plan()
    out = observe(plan, 1, [syn.RingSpec(60.0, 8.0, CENTER, height_fn=syn.flat(6.0))])
    g = out.geometry
    assert g.bead_width_mean_mm == pytest.approx(8.0, rel=0.25)
    assert g.bead_width_bins == 36


def test_bead_width_profile_on_an_ideal_annulus():
    from tasni.modules.extrusion.processing import bead_width_profile
    rng = np.random.default_rng(1)
    theta = rng.uniform(0, 2 * np.pi, 20000)
    r = rng.uniform(36.0, 44.0, 20000)                # annulus 40 +/- 4 -> width 8
    pts = np.column_stack((r * np.cos(theta), r * np.sin(theta), np.zeros(20000)))
    w = bead_width_profile(pts, (0.0, 0.0), bins=36)
    assert w["bins_with_data"] == 36
    assert w["mean_mm"] == pytest.approx(8.0, abs=0.6)   # p97.5 - p2.5 of a uniform 8 mm band


# ------------------------------------------------- characterize a ring (Task 6)

from tasni.modules.extrusion.processing import characterize_ring, fit_circle_xy


def test_characterize_recovers_a_ring_the_recipe_got_wrong():
    pytest.importorskip("open3d")
    # The recipe/plan says 75 mm radius, 6 mm bead. The physical ring is 60 / 8,
    # 6 mm tall, and sits 15 mm off the table centre.
    plan = scene_plan(radius=75.0, bead=6.0, layer_height=5.0)
    true_center = (CENTER[0] + 15.0, CENTER[1] - 10.0)
    T = syn.inspection_camera_T(aim_point_mm(plan.recipe, plan.setup, 1), 300.0)
    depth = syn.render_scene([syn.RingSpec(60.0, 8.0, true_center, height_fn=syn.flat(6.0))], T,
                             plane_center_xy_mm=CENTER)
    found = characterize_ring(depth=depth,
                              geometry=syn.geometry(),
                              T_work_camera=T,
                              search_center_mm=CENTER, work_frame="Tasni Work Frame",
                              config=ExtrusionConfig())
    assert found.radius_mm == pytest.approx(60.0, abs=1.0)
    assert found.center_mm[0] == pytest.approx(true_center[0], abs=1.0)
    assert found.center_mm[1] == pytest.approx(true_center[1], abs=1.0)
    assert found.bead_width_mm == pytest.approx(8.0, abs=2.0)
    assert found.top_z_mean_mm == pytest.approx(6.0, abs=1.5)
    assert found.report["coarse"]["radius_mm"] == pytest.approx(60.0, abs=3.0)
    assert found.measured_xyz.shape[1] == 3


def test_characterize_selects_ring_instead_of_larger_raised_patch():
    pytest.importorskip("open3d")
    true_center = (CENTER[0] + 8.0, CENTER[1] - 5.0)
    T = syn.inspection_camera_T([CENTER[0], CENTER[1], 6.0], 300.0)
    ring = syn.RingSpec(40.0, 14.0, true_center, height_fn=syn.flat(8.0))
    # A broad 3 mm-high residual, separated from the ring, models the real
    # checkerboard depth bias that used to win solely because it was largest.
    patch_x = np.arange(CENTER[0] + 80.0, CENTER[0] + 151.0, 1.0)
    patch_y = np.arange(CENTER[1] - 110.0, CENTER[1] + 111.0, 1.0)
    X, Y = np.meshgrid(patch_x, patch_y, indexing="ij")
    patch = np.column_stack((X.ravel(), Y.ravel(), np.full(X.size, 3.0)))
    scene = np.vstack((
        syn.plane_points(center_xy_mm=CENTER), ring.surface_points(), patch))
    depth = syn.render_depth(scene, T)
    found = characterize_ring(
        depth=depth, geometry=syn.geometry(),
        T_work_camera=T,
        search_center_mm=CENTER, work_frame="Tasni Work Frame",
        config=ExtrusionConfig())

    assert found.radius_mm == pytest.approx(40.0, abs=1.5)
    assert found.center_mm == pytest.approx(true_center, abs=1.5)
    candidates = found.report["ring_selector"]["candidates"]
    selected = next(candidate for candidate in candidates if candidate.get("selected"))
    assert selected["points"] < max(candidate["points"] for candidate in candidates)
    assert selected["angular_coverage"] >= 0.95
    assert selected["radial_span_ratio"] < 0.8


def test_characterize_real_checkerboard_capture_selects_the_visible_ring():
    """trial 20260828-171615-f088cf48/characterize-01: ring plus board residual.

    The protected behaviour is the SELECTION -- the ring-shape gate must pick
    the ring and not the larger above-plane board residual the old
    largest-cluster rule chose (it reported a 51 mm bead). That is unmoved:
    the selected candidate's coarse circle still reads r 41.13 mm, against
    41.12 under the colour-gate chain.

    The refined numbers below DID move, and were re-measured 2026-08-31 rather
    than carried over. This is a 1 mm-worded pre-protocol-2 frame whose
    substrate fits at -1.36 mm with sigma 0.805, so the derived floor sits
    2.0 mm above THAT -- roughly 0.64 mm in work-frame Z where the old constant
    sat at 2.5 mm. Less of the ring's own low flank is cut away and the board
    no longer arrives as a second large candidate, so the measured bead
    footprint narrows (13.26 -> 9.82 mm) and the crest-read radius moves out
    from 39.17 to 40.39 -- 0.74 mm inside the coarse circle where it used to sit
    1.95 mm inside it. Both moves are the same fact: the old footprint had
    board fused into it, which widened the bead and dragged the crest inward.
    There is no ground truth for this ring in the archive, so what is asserted
    is self-consistency (refined radius near the coarse fit, bead near the
    ring's own width) plus the selection that has not changed.
    """
    pytest.importorskip("open3d")
    fixture = np.load(Path(__file__).parent / "fixtures" / "extrusion" / "ring1"
                      / "ring1_checkerboard_20260828.npz")
    depth = fixture["depth"]
    # This fixture is a pre-protocol-2 capture with 1 mm depth WORDS, so it is
    # processed at the 2 mm voxel it was captured under. The new 1 mm default
    # (spec 4.4) sits at this archive's quantisation floor, where it merges
    # nothing and lets noise through -- it flips this frame's branch-guard
    # outcome in BOTH directions across the two archived takes. On protocol-2
    # depth (0.1 mm words) 1 mm spans ten quantisation steps, which is the point.
    found = characterize_ring(
        depth=depth,
        geometry=gf.aligned(fixture["K"], (depth.shape[1], depth.shape[0])),
        T_work_camera=fixture["T_work_camera"],
        search_center_mm=fixture["search_center_mm"],
        work_frame="Tasni Work Frame", config=ExtrusionConfig(voxel_size_m=0.002))

    assert found.radius_mm == pytest.approx(40.39, abs=0.5)
    assert found.center_mm == pytest.approx((218.56, 150.18), abs=0.5)
    assert found.bead_width_mm == pytest.approx(9.82, abs=0.75)
    assert found.top_z_mean_mm == pytest.approx(6.78, abs=0.75)
    selector = found.report["ring_selector"]
    selected = next(candidate for candidate in selector["candidates"]
                    if candidate.get("selected"))
    # The selection itself is the protected behaviour, and it is UNCHANGED:
    # the coarse circle through the chosen footprint still reads 41.12/41.13 mm.
    assert selected["radius_mm"] == pytest.approx(41.12, abs=0.5)
    assert selected["angular_coverage"] >= 0.95
    assert selected["radial_span_ratio"] <= 0.8
    # The board no longer arrives as a separate, larger candidate at all -- it
    # is below the derived floor rather than above the old constant one -- so
    # "the ring is not the largest blob" is no longer the thing under test here.
    # It stays pinned on the synthetic scene, where the residual is placed on
    # purpose: test_characterize_selects_ring_instead_of_larger_raised_patch.
    largest = max(selector["candidates"], key=lambda candidate: candidate["points"])
    assert selected["points"] == largest["points"]


def _thin_at(mean_mm: float, dips_deg, width_deg: float = 12.0, floor_mm: float = 0.5):
    """A ring that all but vanishes at ``dips_deg`` -- the real low-relief failure.

    The 2026-08-29 capture had a hand-placed dried ring 2-11 mm tall whose two
    thinnest arcs fell under the ROI height floor, so one loop reached DBSCAN as
    two disconnected arcs and the per-cluster shape gate rejected both.

    ``floor_mm`` is the dip's own crest height, and it has to sit under the
    ROI floor for the scene to reproduce that failure at all. It was 1.0,
    chosen against the old CONSTANT 2.5 mm floor; the derived floor on this
    synthetic plane measures 1.50 mm, and 1.0 mm of dip lands just above it
    once the substrate fit's bias correction is applied -- the ring arrives
    whole and there is nothing to assemble. Measured 2026-08-31 at dips
    1.0/0.5/0.2 mm: 1.0 gives one cluster at coverage 1.000, while 0.5 and 0.2
    both give the two arcs (best single 0.611, assembled 0.889) this fixture is
    about. 0.5 it is.
    """
    def height(theta):
        h = np.full_like(theta, float(mean_mm), dtype=float)
        for dip in dips_deg:
            d = np.abs(np.mod(np.degrees(theta) - dip + 180.0, 360.0) - 180.0)
            h = np.where(d <= width_deg, float(floor_mm), h)
        return h
    return height


def test_characterize_assembles_one_ring_from_arcs_the_height_floor_broke():
    pytest.importorskip("open3d")
    # One physical ring, thinned to nothing at 125 deg and 245 deg: the ROI floor
    # erases those arcs, so DBSCAN yields two clusters neither of which spans
    # 70% of the circumference on its own.
    center = (CENTER[0] + 3.0, CENTER[1] - 1.0)
    T = syn.inspection_camera_T([CENTER[0], CENTER[1], 6.0], 300.0)
    ring = syn.RingSpec(40.0, 10.0, center,
                        height_fn=_thin_at(7.0, (125.0, 245.0)))
    depth = syn.render_scene([ring], T, plane_center_xy_mm=CENTER)
    found = characterize_ring(
        depth=depth, geometry=syn.geometry(),
        T_work_camera=T,
        search_center_mm=CENTER, work_frame="Tasni Work Frame",
        config=ExtrusionConfig())

    selector = found.report["ring_selector"]
    selected = next(c for c in selector["candidates"] if c.get("selected"))
    # The winner must be the ASSEMBLED ring, not either arc on its own.
    assert selected["cluster_count"] >= 2
    # 2 x 24 deg of the ring is genuinely erased, so ~0.89 is the honest ceiling;
    # what matters is that it clears the 0.70 gate no single arc could.
    assert selected["angular_coverage"] >= 0.85
    assert max(c["angular_coverage"] for c in selector["candidates"]
               if c["cluster_count"] == 1) < 0.70
    assert found.radius_mm == pytest.approx(40.0, abs=1.5)
    assert found.center_mm == pytest.approx(center, abs=1.5)


def test_characterize_real_low_relief_capture_is_not_rejected():
    """trial 20260829-151445-acb42814/characterize-01: a 2-11 mm dried ring.

    Under the old CONSTANT 2.5 mm floor its two thinnest arcs fell out, so one
    physical loop reached DBSCAN as two disconnected arcs (48/72 and 25/72
    bins), every cluster failed the per-cluster shape gate on its own, and only
    arc assembly rescued the frame.

    The derived floor removes the fragmentation at its source. This capture's
    substrate fits at -1.81 mm with sigma 0.836 mm (a 1 mm-worded pre-protocol-2
    frame), so the floor lands 2.0 mm above THAT -- about 0.2 mm in work-frame Z,
    against 2.5 mm before. The thin arcs clear it, and the ring arrives as ONE
    complete cluster covering 72/72 (measured 2026-08-31: r 42.11, previously
    42.0 via assembly). That is the point of measuring height above the surface
    the deposit rests on rather than above a nominal plane the surface is not on.
    Arc assembly itself stays covered by
    test_characterize_assembles_one_ring_from_arcs_the_height_floor_broke, whose
    synthetic dips are cut below the derived floor on purpose.
    """
    pytest.importorskip("open3d")
    fixture = np.load(Path(__file__).parent / "fixtures" / "extrusion" / "ring1"
                      / "ring1_low_relief_20260829.npz")
    depth = fixture["depth"]
    found = characterize_ring(
        depth=depth,
        geometry=gf.aligned(fixture["K"], (depth.shape[1], depth.shape[0])),
        T_work_camera=fixture["T_work_camera"],
        search_center_mm=fixture["search_center_mm"],
        work_frame="Tasni Work Frame", config=ExtrusionConfig())

    assert found.radius_mm == pytest.approx(42.0, abs=2.0)
    selected = next(c for c in found.report["ring_selector"]["candidates"]
                    if c.get("selected"))
    # No longer assembled, and no longer NEEDS to be: one whole ring, all round.
    assert selected["cluster_count"] == 1
    assert selected["angular_coverage"] == pytest.approx(1.0, abs=0.02)


def test_a_ring_measured_only_in_part_is_not_closed_into_a_full_one():
    pytest.importorskip("open3d")
    # Same broken ring, but through the layer pipeline: the centreline must cover
    # only what was actually measured, and the report must say so.
    center = (CENTER[0], CENTER[1])
    plan = scene_plan(radius=40.0, bead=10.0, layer_height=7.0)
    T = syn.inspection_camera_T(aim_point_mm(plan.recipe, plan.setup, 1), 300.0)
    ring = syn.RingSpec(40.0, 10.0, center, height_fn=_thin_at(7.0, (200.0,), width_deg=35.0))
    depth = syn.render_scene([ring], T, plane_center_xy_mm=CENTER)
    result = process_observation(depth=depth,
                                 geometry=syn.geometry(),
                                 T_work_camera=T,
                                 plan=plan, layer=plan.layers[0],
                                 config=ExtrusionConfig())

    measured = np.asarray(result.measured_xyz, dtype=float)
    fitted, _ = fit_circle_xy(measured)
    theta = np.mod(np.arctan2(measured[:, 1] - fitted[1],
                              measured[:, 0] - fitted[0]), 2 * np.pi)
    occupied = len(set((theta / (2 * np.pi) * 72).astype(int).tolist()))
    assert occupied < 68                       # the gap must survive into the output
    # compare_circle already guards completeness; forcing the spline closed made
    # it blind, because a periodic curve is always 100% complete.
    assert result.report["closed"] is False
    assert result.metrics.path_completeness < 0.95
    assert result.metrics.maximum_angular_gap_deg > 30.0
    assert result.metrics.valid is False
    assert any("completeness" in w or "angular gap" in w for w in result.metrics.warnings)

# ------------------------------------------------------------- archive (Task 7)


from tasni.modules.extrusion.archive import ExtrusionArchive
from tasni.modules.extrusion.models import LayerManifest


def test_archive_keeps_every_take_of_a_layer_and_records_the_mode(tmp_path):
    plan = scene_plan()
    archive = ExtrusionArchive(tmp_path)
    trial = archive.create_trial("t1", plan, mode="MEASURE_ONLY",
                                 experiment={"note": "dried rings, hand-placed"})
    data = json.loads((trial / "trial.json").read_text())
    assert data["mode"] == "MEASURE_ONLY" and data["experiment"]["note"].startswith("dried")
    nominal = np.zeros((4, 3))
    for take in (1, 2):
        manifest = LayerManifest(trial_id="t1", layer_index=2, take=take, mode="MEASURE_ONLY",
                                 recipe=plan.recipe, toolpath_fingerprint=plan.fingerprint,
                                 annotation={"introduced_offset_mm": [10, 0]})
        archive.write_layer(manifest, nominal_xyz=nominal, commanded_xyz=nominal)
    assert (tmp_path / "t1" / "layer-002" / "manifest.json").is_file()
    assert (tmp_path / "t1" / "layer-002-take02" / "manifest.json").is_file()
    assert archive.layer_dir("t1", 2, take=2).name == "layer-002-take02"
    loaded = json.loads((tmp_path / "t1" / "layer-002-take02" / "manifest.json").read_text())
    assert loaded["take"] == 2 and loaded["annotation"]["introduced_offset_mm"] == [10, 0]


def test_archive_writes_a_characterization_directory(tmp_path):
    plan = scene_plan()
    archive = ExtrusionArchive(tmp_path)
    archive.create_trial("t1", plan, mode="MEASURE_ONLY")
    out = archive.write_characterization(
        "t1", 1, color=np.zeros((4, 4, 3), np.uint8), depth=np.zeros((4, 4), np.uint16),
        measured_xyz=np.zeros((5, 3)),
        derived_images={"comparison.png": np.zeros((4, 4, 3), np.uint8)},
        report={"radius_mm": 60.0})
    assert out.name == "characterize-01"
    for name in ("color.png", "depth.npy", "measured_path.json", "comparison.png", "report.json"):
        assert (out / name).is_file(), name


# ------------------------------------------------- MEASURE_ONLY job (Task 8)

from test_extrusion_job import (Ctx, FakeCamera, FakeRdk, START_JOINTS,  # noqa: F401
                                SIDE_APPROACH_JOINTS, SIDE_JOINTS,
                                services)
from tasni.modules.extrusion import measure as measure_mod
from tasni.modules.extrusion import processing as processing_mod
from tasni.modules.extrusion.measure import MeasureSession, RingMeasureJob
from tasni.modules.extrusion.models import DeviationMetrics, RingGeometry
from tasni.modules.extrusion.processing import ProcessingResult


def fake_measure_processing(**kwargs):
    layer = kwargs["layer"]
    pts = np.array([[p.x_mm, p.y_mm, p.z_mm + 6.0] for p in layer.points])
    metrics = DeviationMetrics(mean_absolute_mm=6.4, rms_mm=7.1, maximum_mm=10.0,
                               measured_center_mm=(10.0, 0.0), measured_radius_mm=40.0,
                               path_completeness=0.99, maximum_angular_gap_deg=5, valid=True,
                               center_offset_mm=(10.0, 0.0), center_offset_norm_mm=10.0,
                               shape_rms_mm=0.3, shape_max_mm=0.8)
    geometry = RingGeometry(top_z_mean_mm=6, top_z_min_mm=5, top_z_max_mm=7, top_z_std_mm=0.5,
                            height_mean_mm=6, height_min_mm=5, height_max_mm=7,
                            height_reference="build_plane", bead_width_mean_mm=8,
                            bead_width_min_mm=7, bead_width_max_mm=9, bead_width_bins=36)
    image = np.zeros((12, 12), np.uint8)
    fake_measure_processing.calls.append(kwargs)
    return ProcessingResult(pts, None, metrics, image, image, np.zeros((12, 12, 3), np.uint8),
                            {"counts": {"raw_depth_pixels": 256}, "timings_ms": {"total_ms": 10.0},
                             "branch_guard_attempts": [{"attempt": 1}]},
                            filtered_xyz=pts.copy(), geometry=geometry)


fake_measure_processing.calls = []


def measure_env(tmp_path, monkeypatch, *, hardware_approved=False, side_photo=False):
    svc, rdk, camera = services(tmp_path)
    svc.config.extrusion.hardware_io_test_approved = hardware_approved
    # The side photo is ON in the cell (it is part of the protocol) but OFF for
    # the tests that count grabs and moves: it would add one of each to every
    # assertion about the MEASUREMENT path, which is not what they are about.
    # The tests that own it turn it back on.
    svc.config.extrusion.side_capture_enabled = side_photo
    # Most orchestration tests count logical takes, not the depth-fusion burst.
    # Dedicated tests below turn the production default back on.
    svc.config.extrusion.measure_depth_fusion_frames = 1
    monkeypatch.setattr(measure_mod, "REPO_ROOT", tmp_path)
    monkeypatch.setattr(processing_mod, "process_observation", fake_measure_processing)
    monkeypatch.setattr(measure_mod, "_git_commit", lambda: "abc123")
    monkeypatch.setattr(measure_mod.time, "sleep", lambda _: None)
    monkeypatch.setattr(measure_mod.runs, "read_active", lambda module: {"run_id": "cal-1"})
    fake_measure_processing.calls.clear()
    return svc, rdk, camera


def auto_plan(layers=3):
    recipe = CylinderRecipe(radius_mm=40, layer_count=layers, layer_height_mm=6,
                            bead_diameter_mm=8, robot_speed_mm_s=75, extrusion_rate_pct=0,
                            points_per_circle=24)
    setup = CylinderSetup(print_tool="LongCalibTool", work_frame="Tasni Work Frame",
                          inspection_tool="Realsense", inspection_auto=True,
                          center_x_mm=200, center_y_mm=150)
    return generate_cylinder_plan(recipe, setup)


def test_measure_moves_only_the_camera_and_never_touches_the_valve(tmp_path, monkeypatch):
    svc, rdk, camera = measure_env(tmp_path, monkeypatch, hardware_approved=False)
    plan = auto_plan()
    session = MeasureSession.create(tmp_path / "runs" / "extrusion", plan, note="rings")
    out = RingMeasureJob(svc, plan, session, 1, annotation={"introduced_offset_mm": None},
                         check_collisions=True)(Ctx())
    kinds = [e[0] for e in rdk.events]
    assert "station-program" not in kinds and "create" not in kinds     # no valve, no layer program
    assert "create-target" in kinds and "create-inspection" in kinds
    assert ("start", "TasniCylinder_MEASURE_%s_L001_Inspect" % plan.fingerprint[:10], True) in rdk.events
    assert rdk.events[-1] == ("move-joints", START_JOINTS)
    assert any(name.endswith("_Inspect") for name in rdk.deleted)
    assert camera.grabs == 2                                             # readiness + one measurement
    assert out["kind"] == "ring_measure" and out["mode"] == "MEASURE_ONLY"
    layer_dir = Path(out["layer_dir"])
    assert layer_dir.name == "layer-001"
    manifest = json.loads((layer_dir / "manifest.json").read_text())
    assert manifest["mode"] == "MEASURE_ONLY" and manifest["take"] == 1
    assert manifest["geometry"]["bead_width_mean_mm"] == 8
    timings = manifest["processing"]["timings_ms"]
    assert timings["capture_ms"] >= 0
    assert timings["acquisition_to_path_ms"] == pytest.approx(timings["capture_ms"] + 10.0)
    assert (layer_dir / "depth.npy").is_file() and (layer_dir / "color.png").is_file()


def test_nonzero_depth_median_does_not_turn_missing_pixels_into_near_surfaces():
    from tasni.modules.extrusion.measure import _nonzero_median_depth

    burst = np.array([
        [[0, 100, 100], [0, 100, 120]],
        [[0, 110, 0], [200, 110, 110]],
        [[0, 120, 120], [220, 120, 100]],
    ], dtype=np.uint16)

    assert _nonzero_median_depth(burst).tolist() == [
        [0, 110, 110], [210, 110, 110]]


def test_fusion_geometry_ignores_live_telemetry_but_not_calibration_changes():
    """Temperature is not geometry; an intrinsics change still stops fusion."""
    from dataclasses import replace
    from tasni.modules.extrusion.measure import _same_reconstruction_geometry

    first = gf.aligned(syn.K_720P, (16, 16))
    changed_telemetry = replace(
        first, temps={"projector_c": 48.2},
        raw={**first.raw, "temps": {"projector_c": 48.2}})
    changed_intrinsics = replace(first, depth_K=first.depth_K.copy())
    changed_intrinsics.depth_K[0, 0] += 1.0

    assert _same_reconstruction_geometry(first, changed_telemetry) is True
    assert _same_reconstruction_geometry(first, changed_intrinsics) is False


def test_measure_fuses_five_top_frames_and_archives_the_raw_burst(tmp_path, monkeypatch):
    svc, rdk, camera = measure_env(tmp_path, monkeypatch)
    svc.config.extrusion.measure_depth_fusion_frames = 5
    plan = auto_plan()
    session = MeasureSession.create(tmp_path / "runs" / "extrusion", plan)
    ctx = Ctx()

    out = RingMeasureJob(svc, plan, session, 1, check_collisions=False)(ctx)

    layer = Path(out["layer_dir"])
    raw = np.load(layer / "depth-frames.npy", allow_pickle=False)
    manifest = json.loads((layer / "manifest.json").read_text())
    fusion = manifest["provenance"]["depth_fusion"]
    assert camera.grabs == 6                         # readiness + five-frame burst
    assert camera.streams == 1                       # one greeting/connection for the burst
    assert raw.shape == (5, 16, 16)
    assert fusion["captured_frames"] == 5
    assert fusion["method"] == "per-pixel nonzero median"
    assert manifest["processing"]["depth_plane_check"]["fusion_frames"] == 5
    assert ctx.checkpoints[-1][0] == "extrusion_take"
    assert ctx.checkpoints[-1][1]["take"] == 1


def test_repeat_takes_share_one_session(tmp_path, monkeypatch):
    svc, rdk, camera = measure_env(tmp_path, monkeypatch)
    plan = auto_plan()
    root = tmp_path / "runs" / "extrusion"
    session = MeasureSession.create(root, plan)
    RingMeasureJob(svc, plan, session, 1, annotation={}, check_collisions=True)(Ctx())
    second = RingMeasureJob(svc, plan, session, 1, annotation={"note": "re-placed"},
                            check_collisions=True)(Ctx())
    assert Path(second["layer_dir"]).name == "layer-001-take02"
    assert fake_measure_processing.calls[-1]["assemble_arcs"] is True
    third = RingMeasureJob(svc, plan, session, 2, annotation={"introduced_offset_mm": [10, 0]},
                           check_collisions=True)(Ctx())
    assert fake_measure_processing.calls[-1]["assemble_arcs"] is False
    assert json.loads((Path(third["layer_dir"]) / "manifest.json").read_text())["annotation"] == {"introduced_offset_mm": [10, 0]}
    # Session survives a restart.
    reloaded = MeasureSession.load(root, session.trial_id)
    assert reloaded.takes == {1: 2, 2: 1}
    assert MeasureSession.latest(root).trial_id == session.trial_id
    assert reloaded.last_pose is not None


def test_measure_archives_the_raw_frame_when_processing_fails(tmp_path, monkeypatch):
    svc, rdk, _ = measure_env(tmp_path, monkeypatch)
    monkeypatch.setattr(processing_mod, "process_observation",
                        lambda **kwargs: (_ for _ in ()).throw(RuntimeError("bad skeleton")))
    plan = auto_plan()
    session = MeasureSession.create(tmp_path / "runs" / "extrusion", plan)
    with pytest.raises(RuntimeError, match="raw RGB-D archived"):
        RingMeasureJob(svc, plan, session, 1, annotation={}, check_collisions=True)(Ctx())
    layer = session.trial_dir / "layer-001"
    assert (layer / "depth.npy").is_file() and "bad skeleton" in (layer / "report.json").read_text()
    assert rdk.events[-1] == ("move-joints", START_JOINTS)                     # still returns home


def test_measure_blocks_before_motion_when_the_camera_is_offline(tmp_path, monkeypatch):
    from tasni.core.camera import CameraError
    svc, rdk, camera = measure_env(tmp_path, monkeypatch)
    def offline(**kwargs): raise CameraError("camera timeout (100.123.63.127:1024)")
    camera.grab = offline
    plan = auto_plan()
    session = MeasureSession.create(tmp_path / "runs" / "extrusion", plan)
    with pytest.raises(RuntimeError, match="camera is not ready"):
        RingMeasureJob(svc, plan, session, 1, annotation={}, check_collisions=True)(Ctx())
    assert rdk.events == []


# ---------------------------------------------- characterize job (Task 9)

from tasni.modules.extrusion.measure import RingCharacterizeJob
from tasni.modules.extrusion.processing import CharacterizationResult


def fake_characterize(**kwargs):
    fake_characterize.calls.append(kwargs)
    image = np.zeros((12, 12), np.uint8)
    return CharacterizationResult(
        radius_mm=61.2, center_mm=(214.0, 141.0), bead_width_mm=8.3, bead_width_min_mm=7.0,
        bead_width_max_mm=9.5, top_z_mean_mm=6.4, top_z_min_mm=5.1, top_z_max_mm=9.8,
        measured_xyz=np.zeros((10, 3)), segmentation=image, skeleton=image,
        comparison=np.zeros((12, 12, 3), np.uint8),
        report={"coarse": {"radius_mm": 60.0}, "timings_ms": {"total_ms": 12.0}})


fake_characterize.calls = []


def test_characterize_job_measures_the_ring_and_stores_it_in_the_session(tmp_path, monkeypatch):
    svc, rdk, camera = measure_env(tmp_path, monkeypatch)
    monkeypatch.setattr(measure_mod, "characterize_ring", fake_characterize)
    fake_characterize.calls.clear()
    plan = auto_plan()
    root = tmp_path / "runs" / "extrusion"
    session = MeasureSession.create(root, plan)
    out = RingCharacterizeJob(svc, plan, session, check_collisions=True)(Ctx())
    assert out["kind"] == "ring_characterize"
    assert out["characterization"]["radius_mm"] == 61.2
    assert fake_characterize.calls[-1]["search_center_mm"] == (200.0, 150.0)
    assert fake_characterize.calls[-1]["work_frame"] == "Tasni Work Frame"
    assert Path(out["capture_dir"]).name == "characterize-01"
    assert (Path(out["capture_dir"]) / "depth.npy").is_file()
    assert out["characterization"]["inspection_pose"]["standoff_mm"] == 300.0
    kinds = [e[0] for e in rdk.events]
    assert "station-program" not in kinds and "create" not in kinds
    assert rdk.events[-1] == ("move-joints", START_JOINTS)
    assert MeasureSession.load(root, session.trial_id).characterizations[-1]["radius_mm"] == 61.2


def test_measure_only_close_range_requires_explicit_job_option(tmp_path, monkeypatch):
    svc, _, _ = measure_env(tmp_path, monkeypatch)
    monkeypatch.setattr(measure_mod, "characterize_ring", fake_characterize)
    plan = auto_plan()
    session = MeasureSession.create(tmp_path / "runs" / "extrusion", plan)

    # The shipped default sits at the 1280x720 MinZ-bounded 300 mm floor; use a
    # distinct value here so the test proves the clamp is wired, not a coincidence.
    svc.config.extrusion.measure_close_range_min_mm = 200.0
    out = RingCharacterizeJob(
        svc, plan, session, check_collisions=True,
        close_range_tool_clear=True)(Ctx())

    pose = out["characterization"]["inspection_pose"]
    assert pose["near_mm"] == 200.0
    assert pose["standoff_mm"] == 200.0
    assert pose["d_fit_mm"] < pose["standoff_mm"]
    assert pose["fill_fraction"]["height"] > 0.5


def test_characterize_archives_raw_frame_when_ring_processing_fails(tmp_path, monkeypatch):
    svc, rdk, _ = measure_env(tmp_path, monkeypatch)
    plan = auto_plan()
    root = tmp_path / "runs" / "extrusion"
    session = MeasureSession.create(root, plan)
    # A prior successful characterization occupies index 1; a failed capture
    # must advance to index 2 rather than overwrite it or disappear.
    monkeypatch.setattr(measure_mod, "characterize_ring", fake_characterize)
    RingCharacterizeJob(svc, plan, session, check_collisions=True)(Ctx())
    monkeypatch.setattr(
        measure_mod, "characterize_ring",
        lambda **kwargs: (_ for _ in ()).throw(RuntimeError("no ring-like cluster")))

    with pytest.raises(RuntimeError, match="raw RGB-D archived"):
        RingCharacterizeJob(svc, plan, session, check_collisions=True)(Ctx())

    failed = session.trial_dir / "characterize-02"
    assert (failed / "color.png").is_file()
    assert (failed / "depth.npy").is_file()
    report = json.loads((failed / "report.json").read_text())
    assert not report["valid"] and "no ring-like cluster" in report["error"]
    assert report["depth_shape"] == [16, 16]
    assert rdk.events[-1] == ("move-joints", START_JOINTS)


# ------------------------------------------------------------- API (Task 10)

from fastapi.testclient import TestClient
from tasni.core.config import AppConfig
from tasni.modules.extrusion import module as extrusion_module
from tasni.webapp.server import create_app


def api_plan(client):
    payload = {"recipe": auto_plan().recipe.model_dump(mode="json"),
               "setup": auto_plan().setup.model_dump(mode="json")}
    return client.post("/api/modules/extrusion/generate", json=payload).json()


def test_measure_layer_is_gated_on_fingerprint_confirm_and_connection_only(tmp_path, monkeypatch):
    monkeypatch.setattr(extrusion_module, "REPO_ROOT", tmp_path)
    cfg = AppConfig()
    cfg.extrusion.hardware_io_test_approved = False          # irrelevant to measuring
    client = TestClient(create_app(cfg))
    plan = api_plan(client)
    body = {"fingerprint": "stale", "layer_index": 1, "annotation": {},
            "confirm_robot_motion": True}
    assert client.post("/api/modules/extrusion/measure/layer", json=body).status_code == 409
    body["fingerprint"] = plan["fingerprint"]
    body["confirm_robot_motion"] = False
    assert client.post("/api/modules/extrusion/measure/layer", json=body).status_code == 400
    body["confirm_robot_motion"] = True
    body["layer_index"] = 99
    assert client.post("/api/modules/extrusion/measure/layer", json=body).status_code == 400
    body["layer_index"] = 1
    refused = client.post("/api/modules/extrusion/measure/layer", json=body)
    assert refused.status_code == 409 and "RoboDK" in refused.json()["detail"]
    assert "hardware" not in refused.json()["detail"].lower()


def test_measure_session_is_created_listed_and_excluded_from_print_counts(tmp_path, monkeypatch):
    monkeypatch.setattr(extrusion_module, "REPO_ROOT", tmp_path)
    client = TestClient(create_app(AppConfig()))
    assert client.get("/api/modules/extrusion/measure/session").json()["session"] is None
    assert client.post("/api/modules/extrusion/measure/session/new",
                       json={"note": "x"}).status_code == 409      # needs a generated plan
    api_plan(client)
    created = client.post("/api/modules/extrusion/measure/session/new", json={"note": "rings"}).json()
    trial_id = created["session"]["trial_id"]
    assert (tmp_path / "runs" / "extrusion" / trial_id / "session.json").is_file()
    assert client.get("/api/modules/extrusion/measure/session").json()["session"]["trial_id"] == trial_id
    assert client.get("/api/modules/extrusion/status").json()["measure_session"] == trial_id
    # A LIVE_PRINT trial beside it: only that one is a printed trial.
    live = ExtrusionArchive(tmp_path / "runs" / "extrusion")
    live.create_trial("20990101-000000-live0000", auto_plan())
    trials = client.get("/api/modules/extrusion/trials").json()
    assert trials["summary"]["total_trials"] == 1
    assert trials["summary"]["measure_only_trials"] == 1
    assert {t["trial_id"]: t["mode"] for t in trials["trials"]} == {
        trial_id: "MEASURE_ONLY", "20990101-000000-live0000": "LIVE_PRINT"}


def test_apply_characterization_rewrites_recipe_and_placement(tmp_path, monkeypatch):
    monkeypatch.setattr(extrusion_module, "REPO_ROOT", tmp_path)
    client = TestClient(create_app(AppConfig()))
    before = api_plan(client)
    assert client.post("/api/modules/extrusion/measure/apply-characterization").status_code == 409
    client.post("/api/modules/extrusion/measure/session/new", json={"note": ""})
    session = MeasureSession.latest(tmp_path / "runs" / "extrusion")
    session.characterizations.append({"index": 1, "radius_mm": 61.24, "center_mm": [214.0, 141.0],
                                      "bead_width_mm": 8.31, "top_z_mean_mm": 6.44,
                                      "top_z_min_mm": 5.1, "top_z_max_mm": 9.8})
    session.save()
    after = client.post("/api/modules/extrusion/measure/apply-characterization").json()
    assert after["fingerprint"] != before["fingerprint"]
    assert after["recipe"]["radius_mm"] == 61.2 and after["recipe"]["bead_diameter_mm"] == 8.3
    assert after["recipe"]["layer_height_mm"] == 6.4
    assert after["setup"]["center_x_mm"] == 214.0 and after["setup"]["center_y_mm"] == 141.0
    assert after["setup"]["build_plane_z_mm"] == 0.0
    assert client.get("/api/modules/extrusion/plan").json()["fingerprint"] == after["fingerprint"]


# ------------------------------- offline reprocessing measures against the TAKE's plan

def test_reprocessing_uses_the_recipe_and_centre_the_take_was_measured_against(tmp_path):
    """A measure-only session is created BEFORE Characterize -> Apply, so trial.json
    carries the pre-Apply plan (cell 2026-08-28: r 40 at (212.1, 149.7) while the
    take was measured at r 42.6 about (214.6, 146.7)). Reprocessing from trial.json
    would score the ring against a plan it was never measured against -- the same
    stale-plan artifact the handoff says must never reach the paper.
    """
    pytest.importorskip("open3d")
    from tasni.modules.extrusion.service import reprocess_saved_layer

    root = tmp_path / "runs" / "extrusion"
    stale = scene_plan(radius=40.0, bead=15.0, layer_height=5.0, center=(212.1, 149.7))
    applied = scene_plan(radius=60.0, bead=8.0, layer_height=6.0, center=CENTER)
    archive = ExtrusionArchive(root)
    # As MeasureSession.create writes it: the trial carries NO processing
    # provenance -- the measure job puts intrinsics and the processing config on
    # each take's manifest, because they are per-take facts.
    archive.create_trial("t-stale", stale, mode="MEASURE_ONLY")
    layer = applied.layers[0]
    T = syn.inspection_camera_T(aim_point_mm(applied.recipe, applied.setup, 1), 300.0)
    depth = syn.render_scene([syn.RingSpec(60.0, 8.0, CENTER, height_fn=syn.flat(6.0))], T,
                             plane_center_xy_mm=CENTER, seed=3)
    color = np.zeros((*depth.shape, 3), np.uint8)
    manifest = LayerManifest(
        trial_id="t-stale", layer_index=1, take=1, mode="MEASURE_ONLY",
        recipe=applied.recipe, toolpath_fingerprint=applied.fingerprint,
        color_file="color.png", depth_file="depth.npy",
        processing={"valid": False, "error": "branch guard exhausted"},
        provenance={"T_work_camera": np.asarray(T, dtype=float).tolist(),
                    "camera_intrinsics": {"K": syn.K_720P.tolist()},
                    # The frame's own greeting, exactly as the measure job
                    # records it. Without it the reprocess path falls back to
                    # the legacy 1 mm depth convention and reads this take's
                    # 0.1 mm words 10x too far away.
                    "camera_geometry": syn.geometry_dict(),
                    "processing_config": ExtrusionConfig().model_dump(mode="json")})
    nominal = points_array(layer)
    archive.write_layer(manifest, nominal_xyz=nominal, commanded_xyz=nominal,
                        color=color, depth=depth, report={"valid": False})

    result = reprocess_saved_layer(root, "t-stale", 1)

    metrics = result["metrics"]
    assert metrics["valid"], metrics["warnings"]
    assert metrics["measured_radius_mm"] == pytest.approx(60.0, abs=1.0)
    assert metrics["center_offset_norm_mm"] < 1.0, (
        "scored against the stale trial.json plan instead of the take's own nominal")
    rewritten = json.loads((root / "t-stale" / "layer-001" / "manifest.json").read_text(
        encoding="utf-8"))
    assert rewritten["metrics"]["valid"] is True
    assert rewritten["processing"]["offline_reprocess"] is True
    assert rewritten["recipe"]["radius_mm"] == 60.0, "the take's recipe must survive"
    # paper_summary reads height/bead from manifest.geometry: a reprocessed take
    # that leaves it empty silently drops out of those statistics.
    assert rewritten["geometry"] is not None
    assert rewritten["geometry"]["height_mean_mm"] == pytest.approx(6.0, abs=1.5)
    assert rewritten["geometry"]["bead_width_mean_mm"] > 0


# --------------------------------------------------- paper summary (Task 11)

from tasni.modules.extrusion.measure import paper_summary


def _write_take(root, trial_id, layer_index, take, *, offset, rms, mean_abs, maximum,
                acq_ms, valid=True, offset_norm=None, measured_offset=None,
                phase=None, offline=False, cycle_ms=None):
    if measured_offset is None:
        measured_offset = (float(offset_norm or 0.0), 0.0)
    if offset_norm is None:
        offset_norm = float(np.hypot(*measured_offset))
    # As the chain reports it: the fitted centre IS the plan centre plus the offset.
    setup = auto_plan().setup
    measured_center = (setup.center_x_mm + float(measured_offset[0]),
                       setup.center_y_mm + float(measured_offset[1]))
    annotation = {"introduced_offset_mm": offset}
    if phase is not None:
        annotation["phase"] = phase
    manifest = LayerManifest(
        trial_id=trial_id, layer_index=layer_index, take=take, mode="MEASURE_ONLY",
        recipe=auto_plan().recipe, toolpath_fingerprint="f" * 64,
        annotation=annotation,
        metrics=DeviationMetrics(mean_absolute_mm=mean_abs, rms_mm=rms, maximum_mm=maximum,
                                 measured_center_mm=measured_center, measured_radius_mm=40,
                                 path_completeness=0.99, maximum_angular_gap_deg=4, valid=valid,
                                 center_offset_mm=tuple(float(v) for v in measured_offset),
                                 center_offset_norm_mm=offset_norm, shape_rms_mm=0.4),
        geometry=RingGeometry(top_z_mean_mm=6, top_z_min_mm=5, top_z_max_mm=9, top_z_std_mm=1,
                              height_mean_mm=6, height_min_mm=5, height_max_mm=9,
                              height_reference="build_plane", bead_width_mean_mm=8,
                              bead_width_min_mm=7, bead_width_max_mm=9, bead_width_bins=36),
        processing={"offline_reprocess": offline,
                    "timings_ms": ({"capture_ms": 40.0, "total_ms": acq_ms - 40.0}
                                   if offline else
                                   {"capture_ms": 40.0, "total_ms": acq_ms - 40.0,
                                    "acquisition_to_path_ms": acq_ms,
                                    **({} if cycle_ms is None
                                       else {"inspection_cycle_ms": float(cycle_ms)})})})
    ExtrusionArchive(root).write_layer(manifest, nominal_xyz=np.zeros((4, 3)),
                                       commanded_xyz=np.zeros((4, 3)))


def test_paper_summary_groups_by_introduced_offset_and_reports_timing(tmp_path):
    root = tmp_path / "runs" / "extrusion"
    session = MeasureSession.create(root, auto_plan(), note="rings")
    t = session.trial_id
    _write_take(root, t, 1, 1, offset=None, offset_norm=0.4, rms=0.5, mean_abs=0.4, maximum=1.1, acq_ms=900)
    _write_take(root, t, 1, 2, offset=[0, 0], offset_norm=0.6, rms=0.6, mean_abs=0.5, maximum=1.3, acq_ms=1100)
    _write_take(root, t, 2, 1, offset=[10, 0], offset_norm=9.8, rms=7.0, mean_abs=6.3, maximum=9.9, acq_ms=1000)
    summary = paper_summary(root, t)
    assert summary["mode"] == "MEASURE_ONLY" and summary["takes"] == 3 and summary["valid"] == 3
    by_name = {c["condition"]: c for c in summary["conditions"]}
    # A condition is the layer AND the ground truth: see
    # test_conditions_separate_the_layer_and_the_phase_the_operator_recorded.
    assert by_name["layer 1 - no introduced offset"]["takes"] == 2
    shifted = by_name["layer 2 - introduced offset (10, 0) mm"]
    assert shifted["takes"] == 1 and shifted["center_offset_norm_mm"]["mean"] == 9.8
    assert summary["timing_ms"]["acquisition_to_path_ms"]["mean"] == pytest.approx(1000.0)
    assert summary["timing_ms"]["acquisition_to_path_ms"]["sd"] == pytest.approx(100.0)
    assert summary["height_mm"]["height_max_mm"]["mean"] == 9.0
    assert "10" in summary["markdown"] and "hand-placed" in summary["markdown"]


def test_paper_summary_endpoint(tmp_path, monkeypatch):
    monkeypatch.setattr(extrusion_module, "REPO_ROOT", tmp_path)
    client = TestClient(create_app(AppConfig()))
    api_plan(client)
    trial_id = client.post("/api/modules/extrusion/measure/session/new", json={"note": ""}).json()["session"]["trial_id"]
    _write_take(tmp_path / "runs" / "extrusion", trial_id, 1, 1, offset=None, offset_norm=0.5,
                rms=0.5, mean_abs=0.4, maximum=1.0, acq_ms=950)
    got = client.get(f"/api/modules/extrusion/trials/{trial_id}/paper-summary").json()
    assert got["takes"] == 1 and "markdown" in got
    assert client.get("/api/modules/extrusion/trials/nope/paper-summary").status_code == 404


# ---------------------------------- detection error vs the operator's ground truth

def test_paper_summary_scores_the_measurement_against_the_offset_the_operator_typed(tmp_path):
    """The paper's claim is "a 10 mm shift READ AS 10.x mm", not "offset 10.x mm".

    Without the typed offset subtracted, the table only says where the ring was,
    never how well the chain found it -- which is the whole controlled-validation
    claim.
    """
    root = tmp_path / "runs" / "extrusion"
    t = MeasureSession.create(root, auto_plan(), note="rings").trial_id
    _write_take(root, t, 2, 1, offset=[10, 0], measured_offset=[10.4, 0.3],
                rms=7.0, mean_abs=6.3, maximum=9.9, acq_ms=1000)
    _write_take(root, t, 2, 2, offset=[10, 0], measured_offset=[9.6, -0.3],
                rms=7.1, mean_abs=6.4, maximum=10.1, acq_ms=1000)

    shifted = {c["condition"]: c for c in paper_summary(root, t)["conditions"]}[
        "layer 2 - introduced offset (10, 0) mm"]

    assert shifted["introduced_norm_mm"] == pytest.approx(10.0)
    # |(10.4, 0.3) - (10, 0)| and |(9.6, -0.3) - (10, 0)| are both exactly 0.5 mm
    assert shifted["detection_error_mm"]["mean"] == pytest.approx(0.5, abs=1e-9)
    assert shifted["detection_error_mm"]["max"] == pytest.approx(0.5, abs=1e-9)
    assert shifted["detection_error_mm"]["n"] == 2


def test_a_take_with_no_introduced_offset_is_scored_against_zero(tmp_path):
    """The zero-offset group IS the baseline: how far off it reads when nothing moved."""
    root = tmp_path / "runs" / "extrusion"
    t = MeasureSession.create(root, auto_plan(), note="rings").trial_id
    _write_take(root, t, 1, 1, offset=None, measured_offset=[0.3, 0.4],
                rms=0.5, mean_abs=0.4, maximum=1.1, acq_ms=900)

    baseline = paper_summary(root, t)["conditions"][0]

    assert baseline["introduced_norm_mm"] == 0.0
    assert baseline["detection_error_mm"]["mean"] == pytest.approx(0.5, abs=1e-9)
    assert baseline["shift_consistency"] is None, "no shift to check the relation against"


def test_deviations_that_look_like_the_introduced_shift_pass_the_built_in_relation(tmp_path):
    """A pure translation d must read max = d, mean = 2d/pi, RMS = d/sqrt(2)."""
    root = tmp_path / "runs" / "extrusion"
    t = MeasureSession.create(root, auto_plan(), note="rings").trial_id
    _write_take(root, t, 2, 1, offset=[10, 0], measured_offset=[10.0, 0.0],
                rms=7.07, mean_abs=6.37, maximum=10.0, acq_ms=1000)

    check = paper_summary(root, t)["conditions"][0]["shift_consistency"]

    assert check["consistent"] is True
    assert check["expected_mm"]["mean_absolute_mm"] == pytest.approx(6.366, abs=1e-3)
    assert check["expected_mm"]["rms_mm"] == pytest.approx(7.071, abs=1e-3)
    assert check["disagreements"] == []


def test_deviations_that_do_not_look_like_a_pure_shift_are_flagged_not_averaged(tmp_path):
    """The sanity check the handoff calls "stop and investigate" must be machine-checked."""
    root = tmp_path / "runs" / "extrusion"
    t = MeasureSession.create(root, auto_plan(), note="rings").trial_id
    _write_take(root, t, 2, 1, offset=[10, 0], measured_offset=[3.0, 0.0],
                rms=2.4, mean_abs=2.0, maximum=4.0, acq_ms=1000)

    summary = paper_summary(root, t)
    check = summary["conditions"][0]["shift_consistency"]

    assert check["consistent"] is False
    assert set(check["disagreements"]) == {"mean_absolute_mm", "rms_mm", "maximum_mm"}
    assert "does not match a pure" in summary["markdown"]


def test_the_markdown_states_the_recovered_shift_for_every_offset_condition(tmp_path):
    root = tmp_path / "runs" / "extrusion"
    t = MeasureSession.create(root, auto_plan(), note="rings").trial_id
    _write_take(root, t, 2, 1, offset=[10, 0], measured_offset=[10.2, 0.0],
                rms=7.1, mean_abs=6.4, maximum=10.2, acq_ms=1000)
    _write_take(root, t, 2, 2, offset=[10, 0], measured_offset=[9.8, 0.0],
                rms=7.0, mean_abs=6.3, maximum=9.8, acq_ms=1000)

    markdown = paper_summary(root, t)["markdown"]

    assert "10.0 mm introduced offset was recovered as 10.00 +/- 0.28 mm" in markdown
    assert "detection error" in markdown


def test_a_take_archived_before_the_offset_vector_existed_is_left_out_of_the_score(tmp_path):
    """Missing is not zero: an old manifest must not read as a perfect measurement."""
    root = tmp_path / "runs" / "extrusion"
    t = MeasureSession.create(root, auto_plan(), note="rings").trial_id
    _write_take(root, t, 2, 1, offset=[10, 0], measured_offset=[10.4, 0.3],
                rms=7.0, mean_abs=6.3, maximum=9.9, acq_ms=1000)
    stale = root / t / "layer-002" / "manifest.json"
    payload = json.loads(stale.read_text(encoding="utf-8"))
    payload["metrics"].pop("center_offset_mm")
    stale.write_text(json.dumps(payload), encoding="utf-8")

    condition = paper_summary(root, t)["conditions"][0]

    assert condition["detection_error_mm"]["n"] == 0
    assert condition["detection_error_mm"]["mean"] is None


def test_empty_roi_error_reports_which_band_rejected_the_points():
    """A bare "not enough points" cannot be acted on in the cell. The message must
    say how many points each band admitted and what the observed spread was, so an
    operator can tell a Z-offset (wrong build plane) from a radial miss (wrong
    centre) without re-running the print."""
    pytest.importorskip("open3d")
    plan = scene_plan()
    # A real ring, but sitting 250 mm away from the configured centre: the radial
    # band rejects everything while the height band is perfectly happy.
    far = (CENTER[0] + 250.0, CENTER[1])
    with pytest.raises(RuntimeError) as excinfo:
        observe(plan, 1, [syn.RingSpec(60.0, 8.0, far, height_fn=syn.flat(6.0))])
    msg = str(excinfo.value)
    # WHICH stage refuses is not the point and is allowed to move: a few
    # scattered substrate specks can clear the derived floor, so the ROI is
    # rarely literally empty and the refusal now usually comes one stage later,
    # at DBSCAN. What must never move is the diagnosis reaching the operator --
    # the per-band tallies that separate a wrong centre from a wrong height
    # reference. Both messages carry them.
    assert ("not enough deposited-geometry points" in msg
            or "no deposited cluster survived" in msg), msg
    assert "height" in msg and "radial" in msg          # both bands named
    assert "in_height_band" in msg and "in_radial_band" in msg


def test_measure_only_requests_default_to_collisions_off():
    """The ring stack is not modelled in the station, so measure-only camera moves
    ship with RoboDK collision validation off; the print paths are unaffected."""
    from tasni.modules.extrusion.module import CharacterizeBody, MeasureLayerBody

    assert CharacterizeBody().collision_check_enabled is False
    assert MeasureLayerBody(fingerprint="f", layer_index=1).collision_check_enabled is False


def test_live_measurement_defers_publication_figures(tmp_path, monkeypatch):
    """Matplotlib must never hold the robot at the inspection pose for minutes."""
    svc, rdk, camera = measure_env(tmp_path, monkeypatch)
    plan = auto_plan()
    session = MeasureSession.create(tmp_path / "runs" / "extrusion", plan)

    out = RingMeasureJob(svc, plan, session, 1, check_collisions=False)(Ctx())

    assert not (Path(out["layer_dir"]) / "figures").exists()
    assert rdk.events[-1] == ("move-joints", START_JOINTS)


def archived_take(root, monkeypatch):
    """One real MEASURE_ONLY take on disk, with no figures rendered yet."""
    import test_extrusion_figures as figs
    monkeypatch.setattr(extrusion_module, "REPO_ROOT", root.parent.parent)
    return figs.write_take(root)


def test_layer_figures_are_served_and_rendered_on_first_request(tmp_path, monkeypatch):
    """Takes archived before figures existed must still produce them, robot-free."""
    layer_dir = archived_take(tmp_path / "runs" / "extrusion", monkeypatch)
    client = TestClient(create_app(AppConfig()))
    base = "/api/modules/extrusion/trials/t1/layers/layer-001"
    assert not (layer_dir / "figures").exists()

    found = client.get(f"{base}/figures/plan.png")

    assert found.status_code == 200
    assert found.headers["content-type"] == "image/png"
    assert found.content[:8] == b"\x89PNG\r\n\x1a\n"
    assert (layer_dir / "figures" / "plan.png").is_file()
    assert client.get(f"{base}/figures/profile.pdf").status_code == 200


def test_archived_frames_are_served_from_the_allowlist_only(tmp_path, monkeypatch):
    archived_take(tmp_path / "runs" / "extrusion", monkeypatch)
    client = TestClient(create_app(AppConfig()))
    base = "/api/modules/extrusion/trials/t1/layers/layer-001"

    assert client.get(f"{base}/files/depth.npy").status_code == 404
    assert client.get(f"{base}/files/manifest.json").status_code == 404
    assert client.get(f"{base}/figures/plan.svg").status_code == 404


def test_a_layer_directory_outside_the_trial_is_refused(tmp_path, monkeypatch):
    """The directory name is a URL segment, so it must not be able to escape.

    A literal ``..`` is normalised away by the client and never reaches the
    handler; a percent-encoded one does, which is the case worth pinning.
    """
    archived_take(tmp_path / "runs" / "extrusion", monkeypatch)
    client = TestClient(create_app(AppConfig()))
    base = "/api/modules/extrusion/trials/t1/layers"

    for segment in ("%2e%2e", "%2e", "characterize-01", "figures"):
        found = client.get(f"{base}/{segment}/figures/plan.png")
        assert found.status_code == 404, f"{segment} was served"
    for trial in ("%2e%2e", "%2e"):
        found = client.get(f"/api/modules/extrusion/trials/{trial}"
                           f"/layers/layer-001/figures/plan.png")
        assert found.status_code == 404, f"trial {trial} was served"
    # A separator inside a segment is normalised by the router into some other
    # path entirely; whatever that lands on, it is never an archived file.
    for escaping in (f"{base}/../figures/plan.png",
                     f"{base}/layer-001%2f..%2f../figures/plan.png",
                     "/api/modules/extrusion/trials/t1%2f..%2f../layers/layer-001"
                     "/figures/plan.png"):
        found = client.get(escaping)
        assert "image" not in found.headers["content-type"], f"{escaping} served an image"
        assert "pdf" not in found.headers["content-type"], f"{escaping} served a pdf"


def test_trials_reports_which_takes_have_figures(tmp_path, monkeypatch):
    layer_dir = archived_take(tmp_path / "runs" / "extrusion", monkeypatch)
    client = TestClient(create_app(AppConfig()))

    layer = client.get("/api/modules/extrusion/trials").json()["trials"][0]["layers"][0]
    assert layer["layer_dir"] == "layer-001"
    assert layer["has_figures"] is False

    client.get("/api/modules/extrusion/trials/t1/layers/layer-001/figures/plan.png")
    layer = client.get("/api/modules/extrusion/trials").json()["trials"][0]["layers"][0]
    assert layer["has_figures"] is True


def test_the_trial_stack_figure_is_served_for_the_whole_session(tmp_path, monkeypatch):
    """The stack across layers is the picture that shows an introduced offset."""
    import test_extrusion_figures as figs
    root = tmp_path / "runs" / "extrusion"
    monkeypatch.setattr(extrusion_module, "REPO_ROOT", tmp_path)
    figs.write_take(root, layer_index=1)
    figs.write_take(root, layer_index=2, measured=figs._ring_xyz(z=12.0))
    client = TestClient(create_app(AppConfig()))

    found = client.get("/api/modules/extrusion/trials/t1/figures/stack.png")

    assert found.status_code == 200
    assert found.headers["content-type"] == "image/png"
    assert (root / "t1" / "figures" / "stack.png").is_file()
    assert client.get("/api/modules/extrusion/trials/t1/figures/plan.png").status_code == 404
    assert client.get("/api/modules/extrusion/trials/nope/figures/stack.png").status_code == 404


# ============================================================================
# The operator journey: a session that stays truthful across restarts,
# reprocessing, and the conditions the paper protocol deliberately creates.
# ============================================================================

def _archive_failed_take(root, session, plan, *, capture_ms=2874.0, layer_index=1, take=1):
    """A take that failed on the cell: raw RGB-D archived, nothing derived.

    This is exactly what the paper trial's layer-001 looked like the night the
    board-noise defect was found.
    """
    layer = plan.layers[layer_index - 1]
    T = syn.inspection_camera_T(aim_point_mm(plan.recipe, plan.setup, layer_index), 300.0)
    depth = syn.render_scene([syn.RingSpec(60.0, 8.0, CENTER, height_fn=syn.flat(6.0))], T,
                             plane_center_xy_mm=CENTER, seed=3)
    manifest = LayerManifest(
        trial_id=session.trial_id, layer_index=layer_index, take=take, mode="MEASURE_ONLY",
        recipe=plan.recipe, toolpath_fingerprint=plan.fingerprint,
        color_file="color.png", depth_file="depth.npy",
        annotation={"introduced_offset_mm": None, "note": "MAIN-TEST-1"},
        processing={"valid": False, "error": "branch guard exhausted",
                    "timings_ms": {"capture_ms": capture_ms}},
        provenance={"T_work_camera": np.asarray(T, dtype=float).tolist(),
                    "camera_intrinsics": {"K": syn.K_720P.tolist()},
                    # The frame's own greeting, exactly as the measure job
                    # records it. Without it the reprocess path falls back to
                    # the legacy 1 mm depth convention and reads this take's
                    # 0.1 mm words 10x too far away.
                    "camera_geometry": syn.geometry_dict(),
                    "processing_config": ExtrusionConfig().model_dump(mode="json")})
    nominal = points_array(layer)
    ExtrusionArchive(root).write_layer(
        manifest, nominal_xyz=nominal, commanded_xyz=nominal,
        color=np.zeros((*depth.shape, 3), np.uint8), depth=depth, report={"valid": False})
    session.takes[layer_index] = take
    session.save()


def test_reprocessing_an_archived_take_puts_it_back_into_the_session(tmp_path):
    """A take rescued offline must re-enter the session, not only its manifest.

    session.json is what the operator's table and the stack view read. On the
    cell (2026-08-28) the paper trial's layer-001 was reprocessed to a valid
    measurement while session.json still said records: [] and tops: {} -- so a
    rescued take was invisible everywhere except its own manifest.
    """
    pytest.importorskip("open3d")
    from tasni.modules.extrusion.service import reprocess_saved_layer

    root = tmp_path / "runs" / "extrusion"
    plan = scene_plan(radius=60.0, bead=8.0, layer_height=6.0, center=CENTER)
    session = MeasureSession.create(root, plan, note="rings")
    _archive_failed_take(root, session, plan)

    reprocess_saved_layer(root, session.trial_id, 1)

    reloaded = MeasureSession.load(root, session.trial_id)
    assert reloaded.takes == {1: 1}
    assert [(r["layer_index"], r["take"]) for r in reloaded.records] == [(1, 1)]
    assert reloaded.records[0]["valid"] is True
    assert reloaded.records[0]["reprocessed"] is True
    assert reloaded.records[0]["layer_name"] == "layer-001"     # the UI addresses figures by it
    assert np.asarray(reloaded.tops[1]).shape[1] == 3           # the stack view can draw it


def test_a_reprocessed_take_keeps_the_capture_time_it_was_measured_with(tmp_path):
    """Capture time is a fact about the cell; offline processing time is not.

    Overwriting capture_ms with an offline number would corrupt the paper's
    acquisition-to-path statistic with a measurement that never happened.
    """
    pytest.importorskip("open3d")
    from tasni.modules.extrusion.service import reprocess_saved_layer

    root = tmp_path / "runs" / "extrusion"
    plan = scene_plan(radius=60.0, bead=8.0, layer_height=6.0, center=CENTER)
    session = MeasureSession.create(root, plan)
    _archive_failed_take(root, session, plan, capture_ms=2874.0)

    reprocess_saved_layer(root, session.trial_id, 1)

    manifest = json.loads((session.trial_dir / "layer-001" / "manifest.json").read_text())
    timings = manifest["processing"]["timings_ms"]
    assert timings["capture_ms"] == pytest.approx(2874.0)
    assert manifest["processing"]["offline_reprocess"] is True
    # No acquisition-to-path: this take never produced a path on the cell.
    assert "acquisition_to_path_ms" not in timings


def test_an_offline_reprocessed_take_is_kept_out_of_the_live_timing_statistic(tmp_path):
    """Requirement 3 is scan-to-feedback time ON THE CELL, over ~10 runs."""
    root = tmp_path / "runs" / "extrusion"
    session = MeasureSession.create(root, auto_plan(), note="rings")
    t = session.trial_id
    _write_take(root, t, 1, 1, offset=None, offset_norm=0.4, rms=0.5, mean_abs=0.4,
                maximum=1.1, acq_ms=1000)
    _write_take(root, t, 1, 2, offset=None, offset_norm=0.5, rms=0.5, mean_abs=0.4,
                maximum=1.2, acq_ms=99999, offline=True)

    summary = paper_summary(root, t)

    acq = summary["timing_ms"]["acquisition_to_path_ms"]
    assert acq["n"] == 1 and acq["mean"] == pytest.approx(1000.0)
    assert summary["timing_ms"]["offline_reprocessed_takes"] == 1
    assert "reprocessed offline" in summary["markdown"]
    # Both takes still count as measurements of the ring.
    assert summary["takes"] == 2


def test_conditions_separate_the_layer_and_the_phase_the_operator_recorded(tmp_path):
    """Five untouched takes and three re-placed takes are different experiments.

    Pooled into one "no introduced offset" row they hide sensing repeatability
    inside placement repeatability, and the per-layer evidence that the camera
    climbs with the stack disappears entirely.
    """
    root = tmp_path / "runs" / "extrusion"
    session = MeasureSession.create(root, auto_plan(), note="rings")
    t = session.trial_id
    for take in (1, 2):
        _write_take(root, t, 1, take, offset=None, offset_norm=0.4, rms=0.5, mean_abs=0.4,
                    maximum=1.1, acq_ms=1000, phase="noise floor")
    _write_take(root, t, 1, 3, offset=None, offset_norm=1.9, rms=1.6, mean_abs=1.4,
                maximum=2.6, acq_ms=1000, phase="re-placed")
    _write_take(root, t, 2, 1, offset=None, offset_norm=0.8, rms=0.7, mean_abs=0.6,
                maximum=1.4, acq_ms=1000, phase="stacked true")
    _write_take(root, t, 2, 2, offset=[10, 0], offset_norm=9.8, rms=7.0, mean_abs=6.3,
                maximum=9.9, acq_ms=1000, phase="top ring shifted")

    summary = paper_summary(root, t)

    by_name = {c["condition"]: c for c in summary["conditions"]}
    assert by_name["layer 1 - noise floor"]["takes"] == 2
    assert by_name["layer 1 - re-placed"]["takes"] == 1
    assert by_name["layer 2 - stacked true"]["takes"] == 1
    shifted = by_name["layer 2 - top ring shifted - introduced offset (10, 0) mm"]
    assert shifted["takes"] == 1 and shifted["layer_index"] == 2
    assert shifted["introduced_norm_mm"] == pytest.approx(10.0)


def test_takes_of_the_same_layer_stay_separate_when_no_phase_was_recorded(tmp_path):
    """Without a phase the layer still separates the conditions."""
    root = tmp_path / "runs" / "extrusion"
    session = MeasureSession.create(root, auto_plan())
    t = session.trial_id
    _write_take(root, t, 1, 1, offset=None, offset_norm=0.4, rms=0.5, mean_abs=0.4,
                maximum=1.1, acq_ms=1000)
    _write_take(root, t, 2, 1, offset=None, offset_norm=0.6, rms=0.6, mean_abs=0.5,
                maximum=1.3, acq_ms=1000)

    by_name = {c["condition"]: c for c in paper_summary(root, t)["conditions"]}

    assert by_name["layer 1 - no introduced offset"]["takes"] == 1
    assert by_name["layer 2 - no introduced offset"]["takes"] == 1


def test_an_invalid_take_is_reported_but_never_averaged_into_the_deviations(tmp_path):
    """An invalid take has no reconstructed path; its numbers are not measurements."""
    root = tmp_path / "runs" / "extrusion"
    session = MeasureSession.create(root, auto_plan())
    t = session.trial_id
    _write_take(root, t, 1, 1, offset=None, offset_norm=0.5, rms=0.5, mean_abs=0.4,
                maximum=1.0, acq_ms=1000, phase="noise floor")
    _write_take(root, t, 1, 2, offset=None, offset_norm=44.0, rms=51.0, mean_abs=48.0,
                maximum=77.0, acq_ms=1000, phase="noise floor", valid=False)

    condition = paper_summary(root, t)["conditions"][0]

    assert condition["takes"] == 2 and condition["valid"] == 1
    assert condition["rms_mm"]["n"] == 1 and condition["rms_mm"]["mean"] == pytest.approx(0.5)
    assert condition["center_offset_norm_mm"]["mean"] == pytest.approx(0.5)


# --------------------------------------------- the plan a session is bound to

def _characterized(client, root, *, radius_mm=61.24, center=(214.0, 141.0)):
    """Characterize + Apply, as the protocol requires between placing and measuring."""
    client.post("/api/modules/extrusion/measure/session/new", json={"note": "rings"})
    session = MeasureSession.latest(root)
    session.characterizations.append({"index": 1, "radius_mm": radius_mm,
                                      "center_mm": list(center), "bead_width_mm": 8.31,
                                      "top_z_mean_mm": 6.44, "top_z_min_mm": 5.1,
                                      "top_z_max_mm": 9.8})
    session.save()
    return client.post("/api/modules/extrusion/measure/apply-characterization").json()


def test_measuring_against_a_plan_that_is_not_the_applied_one_is_refused(tmp_path, monkeypatch):
    """The stale-plan artifact, blocked at the source.

    On the cell (2026-08-28) *Apply to recipe & placement* was skipped between
    Characterize and Measure, and layer-001 measured the ring against a plan it
    was never placed on: 15.38 mm centre offset, 11.31 mm RMS, both meaningless.
    Pressing *Center on scanned surface -> Generate* after Apply does the same
    damage. Once a characterization is applied, a session measures against THAT
    plan or refuses.
    """
    monkeypatch.setattr(extrusion_module, "REPO_ROOT", tmp_path)
    client = TestClient(create_app(AppConfig()))
    api_plan(client)
    applied = _characterized(client, tmp_path / "runs" / "extrusion")

    regenerated = api_plan(client)                       # the pre-Apply plan, again
    assert regenerated["fingerprint"] != applied["fingerprint"]
    refused = client.post("/api/modules/extrusion/measure/layer",
                          json={"fingerprint": regenerated["fingerprint"], "layer_index": 1,
                                "annotation": {}, "confirm_robot_motion": True})

    assert refused.status_code == 409
    detail = refused.json()["detail"]
    assert "Apply to recipe & placement" in detail and "characteriz" in detail.lower()


def test_the_applied_plan_comes_back_after_a_restart(tmp_path, monkeypatch):
    """The plan lives in memory; the session outlives the process.

    "After a backend restart press Apply FIRST" was folklore the operator had to
    remember mid-experiment. The session already records what was applied.
    """
    monkeypatch.setattr(extrusion_module, "REPO_ROOT", tmp_path)
    client = TestClient(create_app(AppConfig()))
    api_plan(client)
    applied = _characterized(client, tmp_path / "runs" / "extrusion")

    restarted = TestClient(create_app(AppConfig()))              # a fresh process
    plan = restarted.get("/api/modules/extrusion/plan").json()

    assert plan["fingerprint"] == applied["fingerprint"]
    assert plan["recipe"]["radius_mm"] == applied["recipe"]["radius_mm"]
    assert plan["setup"]["center_x_mm"] == pytest.approx(214.0)
    assert plan["restored_from"] == MeasureSession.latest(tmp_path / "runs" / "extrusion").trial_id


def test_measuring_layer_two_no_longer_needs_layer_one_measured_first(tmp_path, monkeypatch):
    """The measure-layer-N-1-first gate went with ``floor_profile`` (spec §2.4).

    It existed only because layer N's ROI floor WAS layer N-1's measured top;
    that referencing made the only stacked data worse (completeness 0.62 ->
    0.50), so the data-integrity gate that enforced it is gone too. Layer 2 now
    falls through to the ordinary robot-readiness gate -- still a 409 here (no
    RoboDK in a test client), but no longer one about a missing floor.
    """
    monkeypatch.setattr(extrusion_module, "REPO_ROOT", tmp_path)
    client = TestClient(create_app(AppConfig()))
    api_plan(client)
    applied = _characterized(client, tmp_path / "runs" / "extrusion")

    refused = client.post("/api/modules/extrusion/measure/layer",
                          json={"fingerprint": applied["fingerprint"], "layer_index": 2,
                                "annotation": {}, "confirm_robot_motion": True})

    assert refused.status_code == 409
    detail = refused.json()["detail"]
    assert "measure layer 1 first" not in detail
    assert "connect to RoboDK" in detail


def test_a_failed_take_stays_visible_in_the_session(tmp_path, monkeypatch):
    """A failure the operator cannot see is a failure they cannot reprocess.

    The raw frame is archived, so the take is recoverable -- but only if the
    session shows it happened.
    """
    svc, rdk, _ = measure_env(tmp_path, monkeypatch)
    monkeypatch.setattr(processing_mod, "process_observation",
                        lambda **kwargs: (_ for _ in ()).throw(RuntimeError("bad skeleton")))
    plan = auto_plan()
    session = MeasureSession.create(tmp_path / "runs" / "extrusion", plan)

    with pytest.raises(RuntimeError, match="raw RGB-D archived"):
        RingMeasureJob(svc, plan, session, 1, annotation={"phase": "noise floor"},
                       check_collisions=True)(Ctx())

    reloaded = MeasureSession.load(tmp_path / "runs" / "extrusion", session.trial_id)
    assert len(reloaded.records) == 1
    record = reloaded.records[0]
    assert record["valid"] is False and record["layer_name"] == "layer-001"
    assert "bad skeleton" in record["error"]
    assert record["annotation"] == {"phase": "noise floor"}
    assert 1 not in reloaded.tops                       # a failure is not a measured top


def test_apply_after_a_restart_rebuilds_from_the_sessions_own_trial(tmp_path, monkeypatch):
    """A session that predates the applied-plan record still recovers exactly.

    The paper session (20260828-204846-5b455377) was applied before the session
    recorded what it applied, so a restart can only rebuild it from the trial it
    was created with. Falling back to the module's config defaults instead would
    hand the operator a different work frame, orientation and path resolution --
    a plan the ring was never measured against.
    """
    monkeypatch.setattr(extrusion_module, "REPO_ROOT", tmp_path)
    root = tmp_path / "runs" / "extrusion"
    client = TestClient(create_app(AppConfig()))
    api_plan(client)                                     # points_per_circle 24, not the default 180
    applied = _characterized(client, root)

    session = MeasureSession.latest(root)
    session.applied = None                               # as an archive from before this existed
    session.save()
    restarted = TestClient(create_app(AppConfig()))
    recovered = restarted.post("/api/modules/extrusion/measure/apply-characterization").json()

    assert recovered["fingerprint"] == applied["fingerprint"]
    assert recovered["recipe"]["points_per_circle"] == 24
    assert recovered["setup"]["work_frame"] == "Tasni Work Frame"


# ================================= the Word draft the paper is written from ===

def _docx_content(path):
    """Paragraph texts and tables, as Word will show them."""
    from docx import Document
    document = Document(str(path))
    paragraphs = [p.text for p in document.paragraphs]
    tables = [[[cell.text for cell in row.cells] for row in table.rows]
              for table in document.tables]
    return paragraphs, tables


def _two_condition_trial(root):
    """A noise-floor condition and a 10 mm displaced condition, three takes each."""
    session = MeasureSession.create(root, auto_plan(), note="rings")
    trial_id = session.trial_id
    for take in (1, 2, 3):
        _write_take(root, trial_id, 1, take, offset=None, offset_norm=0.4 + take * 0.1,
                    rms=0.5, mean_abs=0.4, maximum=1.1, acq_ms=1000 + take * 10,
                    phase="noise floor")
    for take in (1, 2, 3):
        _write_take(root, trial_id, 3, take, offset=[10, 0], measured_offset=[10.2, 0.0],
                    rms=7.1, mean_abs=6.4, maximum=10.2, acq_ms=1000,
                    phase="top ring shifted")
    return trial_id


def test_the_word_draft_puts_the_measured_numbers_in_a_real_table(tmp_path):
    """A Markdown block pasted into Word is plain text; the paper needs a table.

    The numbers must also be the SAME object the app reports, never a second
    formatting of the archive that can drift from it.
    """
    pytest.importorskip("docx")
    from tasni.modules.extrusion.paper_docx import build_paper_docx

    root = tmp_path / "runs" / "extrusion"
    trial_id = _two_condition_trial(root)

    out = build_paper_docx(root, trial_id, tmp_path / "draft.docx", embed_figures=False)

    assert out.is_file() and out.suffix == ".docx"
    paragraphs, tables = _docx_content(out)
    # By header, not by position, for the columns too: the draft grows both
    # sections and columns over time, and a positional assertion here just
    # breaks on the next honest addition instead of catching anything.
    conditions = next(t for t in tables if t[0][0] == "Condition")
    column = {name: index for index, name in enumerate(conditions[0])}
    rows = {row[0]: row for row in conditions[1:]}
    assert "layer 1 - noise floor" in rows
    assert rows["layer 1 - noise floor"][column["n"]] == "3"
    shifted = rows["layer 3 - top ring shifted - introduced offset (10, 0) mm"]
    assert shifted[column["n"]] == "3"
    assert "10.20" in shifted[column["Centre offset (mm)"]]
    assert "0.20" in shifted[column["Detection error (mm)"]]
    # How the takes were bought is part of the claim: three frames from one trip
    # and three re-approaches support different sentences about repeatability.
    assert shifted[column["How"]] in {"arm parked", "re-approached", "one take"}
    assert any("10.0 mm introduced offset was recovered" in p for p in paragraphs)


def test_the_draft_says_what_is_still_missing_while_the_run_is_under_way(tmp_path):
    """It is written DURING collection, so it has to show what is not there yet."""
    pytest.importorskip("docx")
    from tasni.modules.extrusion.paper_docx import build_paper_docx

    root = tmp_path / "runs" / "extrusion"
    session = MeasureSession.create(root, auto_plan(), note="rings")
    _write_take(root, session.trial_id, 1, 1, offset=None, offset_norm=0.4, rms=0.5,
                mean_abs=0.4, maximum=1.1, acq_ms=1000, phase="noise floor")

    out = build_paper_docx(root, session.trial_id, tmp_path / "draft.docx",
                           embed_figures=False)

    paragraphs, _ = _docx_content(out)
    text = "\n".join(paragraphs)
    assert "2 more" in text                    # 1 of the 3 takes a condition needs
    assert "11 more" in text                   # 1 of the 12 measurements requirement 3 needs
    assert "not ready to cite" in text.lower()


def test_the_draft_carries_the_wording_the_paper_is_required_to_use(tmp_path):
    """Hand-placed dried beads are not a printed cylinder, and the text says so."""
    pytest.importorskip("docx")
    from tasni.modules.extrusion.paper_docx import build_paper_docx

    root = tmp_path / "runs" / "extrusion"
    trial_id = _two_condition_trial(root)

    out = build_paper_docx(root, trial_id, tmp_path / "draft.docx", embed_figures=False)

    text = "\n".join(_docx_content(out)[0])
    assert "controlled validation" in text
    assert "not the deposition deviation of a printed cylinder" in text
    assert "1.26" in text                      # the error floor, stated where it is claimed


def test_every_take_is_listed_so_the_appendix_needs_no_transcription(tmp_path):
    pytest.importorskip("docx")
    from tasni.modules.extrusion.paper_docx import build_paper_docx

    root = tmp_path / "runs" / "extrusion"
    trial_id = _two_condition_trial(root)

    out = build_paper_docx(root, trial_id, tmp_path / "draft.docx", embed_figures=False)

    takes_table = next(t for t in _docx_content(out)[1] if t[0][0] == "Layer")
    assert takes_table[0][:4] == ["Layer", "Take", "Phase", "Introduced (mm)"]
    assert len(takes_table) == 7                        # header + six takes
    assert [row[0] for row in takes_table[1:]] == ["1", "1", "1", "3", "3", "3"]


def test_the_word_draft_is_served_by_the_api(tmp_path, monkeypatch):
    pytest.importorskip("docx")
    monkeypatch.setattr(extrusion_module, "REPO_ROOT", tmp_path)
    client = TestClient(create_app(AppConfig()))
    api_plan(client)
    trial_id = _two_condition_trial(tmp_path / "runs" / "extrusion")

    got = client.get(f"/api/modules/extrusion/trials/{trial_id}/paper-draft.docx")

    assert got.status_code == 200
    assert got.headers["content-type"] == (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    assert trial_id in got.headers.get("content-disposition", "")
    assert client.get("/api/modules/extrusion/trials/nope/paper-draft.docx").status_code == 404


def test_the_draft_embeds_a_figure_for_every_condition(tmp_path):
    """The figures belong in the draft, not in a folder the writer has to hunt through."""
    pytest.importorskip("docx")
    pytest.importorskip("matplotlib")
    import test_extrusion_figures as figs
    from docx import Document
    from tasni.modules.extrusion.paper_docx import build_paper_docx

    root = tmp_path / "runs" / "extrusion"
    figs.write_take(root, layer_index=1, annotation={"phase": "noise floor"})
    figs.write_take(root, layer_index=2, annotation={"phase": "stacked true"})

    out = build_paper_docx(root, "t1", tmp_path / "draft.docx")

    document = Document(str(out))
    captions = [p.text for p in document.paragraphs if p.text.startswith("Figure ")]
    # The method figure leads, then the stack and the bead as a pipe, then a
    # plan view for each of the two conditions.
    assert len(document.inline_shapes) == len(captions) == 5
    assert "becoming a measured centreline" in captions[0]
    assert any("commanded bead" in c for c in captions)
    assert any("layer 1 - noise floor" in c for c in captions)
    assert any("layer 2 - stacked true" in c for c in captions)


# ------------------------- what inspecting a layer actually costs the print

def test_a_take_records_what_the_inspection_excursion_cost(tmp_path, monkeypatch):
    """The paper has to answer "what does stopping to look cost you?".

    Capture and processing were timed; the robot excursion -- out to the
    viewpoint, settle, and back to the start -- was not, and it is the larger
    half. It is also the only number here that cannot be recovered from the
    archive afterwards, so it has to be taken while the robot is moving.
    """
    svc, rdk, camera = measure_env(tmp_path, monkeypatch)
    plan = auto_plan()
    session = MeasureSession.create(tmp_path / "runs" / "extrusion", plan)

    out = RingMeasureJob(svc, plan, session, 1, annotation={}, check_collisions=True)(Ctx())

    timings = json.loads(
        (Path(out["layer_dir"]) / "manifest.json").read_text())["processing"]["timings_ms"]
    for key in ("move_to_pose_ms", "settle_ms", "capture_ms", "total_ms",
                "return_ms", "inspection_cycle_ms"):
        assert key in timings, f"{key} was not recorded"
    assert timings["settle_ms"] == pytest.approx(svc.config.extrusion.settle_s * 1000.0)
    assert timings["inspection_cycle_ms"] == pytest.approx(
        timings["move_to_pose_ms"] + timings["settle_ms"] + timings["capture_ms"]
        + timings["total_ms"] + timings["return_ms"], abs=1e-6)
    # The session row carries it too, so the operator sees the cost as they go.
    assert "inspection_cycle_ms" in MeasureSession.load(
        tmp_path / "runs" / "extrusion", session.trial_id).records[0]["timings_ms"]


def test_the_summary_reports_what_inspection_costs_per_layer(tmp_path):
    """A duty-cycle number a reviewer can weigh against a layer time."""
    root = tmp_path / "runs" / "extrusion"
    session = MeasureSession.create(root, auto_plan(), note="rings")
    t = session.trial_id
    _write_take(root, t, 1, 1, offset=None, offset_norm=0.4, rms=0.5, mean_abs=0.4,
                maximum=1.1, acq_ms=1000, cycle_ms=14000)
    _write_take(root, t, 1, 2, offset=None, offset_norm=0.5, rms=0.5, mean_abs=0.4,
                maximum=1.2, acq_ms=1000, cycle_ms=16000)

    summary = paper_summary(root, t)

    cycle = summary["timing_ms"]["inspection_cycle_ms"]
    assert cycle["n"] == 2 and cycle["mean"] == pytest.approx(15000.0)
    markdown = summary["markdown"]
    assert "Inspecting one layer cost 15000" in markdown
    assert "rather than during deposition" in markdown


def test_an_unknown_api_path_fails_as_an_api_not_as_the_web_app(tmp_path, monkeypatch):
    """A missing API route must not answer with the single-page app.

    The catch-all that serves client-side routes was answering /api paths too,
    with index.html and status 200. The Word-draft download link saved that web
    page as a .docx -- a 496-byte HTML file named like a document, with no error
    raised anywhere. A route that does not exist has to say so.
    """
    monkeypatch.setattr(extrusion_module, "REPO_ROOT", tmp_path)
    client = TestClient(create_app(AppConfig()))

    missing = client.get("/api/modules/extrusion/there-is-no-such-endpoint")

    assert missing.status_code == 404
    assert "text/html" not in missing.headers.get("content-type", "")


# ------------------ the processing chain, made visible for the method figure

def test_processing_hands_back_every_stage_it_went_through():
    """The method figure has to show what the code did, not a redrawing of it.

    Each stage is handed back as the array the run actually held, and the last
    one IS what was measured -- so a figure built from these cannot drift away
    from the pipeline it claims to illustrate.
    """
    pytest.importorskip("open3d")
    plan = scene_plan(radius=60.0, bead=8.0, layer_height=6.0)
    stages: dict = {}
    result = observe(plan, 1, [syn.RingSpec(60.0, 8.0, CENTER, height_fn=syn.flat(6.0))],
                     stages=stages)

    for key in ("backprojected", "work_roi", "deposit_cluster",
                "radial_trimmed", "top_surface"):
        assert key in stages and len(stages[key]), f"{key} was not handed back"
        assert stages[key].shape[1] == 3
    # Each stage keeps a subset of the one before it.
    assert len(stages["backprojected"]) > len(stages["work_roi"])
    assert len(stages["work_roi"]) >= len(stages["deposit_cluster"])
    assert len(stages["deposit_cluster"]) >= len(stages["top_surface"])
    # The board is in the raw cloud and gone by the ROI.
    assert stages["backprojected"][:, 2].min() < 1.0
    assert stages["work_roi"][:, 2].min() >= 0.5
    assert np.array_equal(stages["top_surface"], result.filtered_xyz)


def test_an_archived_take_rebuilds_the_plan_it_was_measured_against():
    """Shared by offline reprocessing and by the method figure, so they agree."""
    from tasni.modules.extrusion.processing import plan_for_archived_take
    manifest = {"recipe": scene_plan(radius=42.6, bead=12.8).recipe.model_dump(mode="json"),
                "layer_index": 1}
    trial = {"setup": scene_plan().setup.model_dump(mode="json")}
    nominal = points_array(scene_plan(radius=42.6, center=(214.6, 146.7)).layers[0])

    plan = plan_for_archived_take(manifest, trial, nominal_xyz=nominal)

    assert plan.recipe.radius_mm == 42.6
    assert plan.setup.center_x_mm == pytest.approx(214.6, abs=0.01)
    assert plan.setup.center_y_mm == pytest.approx(146.7, abs=0.01)


def _take_with_provenance(root, trial_id, **extra):
    """A take carrying the provenance a real capture records."""
    _write_take(root, trial_id, 1, 1, offset=None, offset_norm=0.4, rms=0.5, mean_abs=0.4,
                maximum=1.1, acq_ms=1000, phase="noise floor", **extra)
    path = root / trial_id / "layer-001" / "manifest.json"
    payload = json.loads(path.read_text(encoding="utf-8"))
    payload["provenance"] = {
        "git_commit": "c48b1ab08a25bb0fa8fb2f6f3b9d3a52f77960e2",
        "camera_resolution": "1280x720",
        "camera_intrinsics": {"K": [[889.87, 0, 648.98], [0, 890.81, 362.0], [0, 0, 1]],
                              "dist_coeffs": [0.1148, -0.2, 0.0, 0.0, 0.0]},
        "calibration": {"run_id": "20260629-130945", "method": "HORAUD",
                        "quality": {"verdict": "borderline", "val_rms_px": 1.115,
                                    "board_consistency_rms_mm": 1.2602}},
        "processing_config": {"deposit_min_height_mm": 0.5, "radial_roi_margin_mm": 30.0,
                              "layer_floor_margin_mm": 2.0, "raster_mm_per_pixel": 1.0,
                              "measured_spline_points": 180, "bead_width_bins": 36,
                              "settle_s": 1.0},
        "inspection_pose": {"standoff_mm": 300.0},
        "work_frame": "Tasni Work Frame", "inspection_tool": "Realsense"}
    payload["processing"]["timings_ms"].update(
        {"move_to_pose_ms": 6100.0, "return_ms": 5400.0, "settle_ms": 1000.0,
         "inspection_cycle_ms": 14460.0})
    path.write_text(json.dumps(payload), encoding="utf-8")


def test_the_draft_states_the_system_it_was_measured_on(tmp_path):
    """A reviewer has to be able to reproduce the setup, not guess at it.

    Every value comes from the take's own provenance -- including the hand-eye
    residual the error-floor claim rests on, which was previously a constant in
    the source and could drift from the calibration actually in use.
    """
    pytest.importorskip("docx")
    from tasni.modules.extrusion.paper_docx import build_paper_docx

    root = tmp_path / "runs" / "extrusion"
    trial_id = MeasureSession.create(root, auto_plan(), note="rings").trial_id
    _take_with_provenance(root, trial_id)

    out = build_paper_docx(root, trial_id, tmp_path / "draft.docx", embed_figures=False)

    paragraphs, tables = _docx_content(out)
    settings = {row[0]: row[1] for table in tables for row in table if len(row) == 2}
    assert "1280x720" in settings["Camera"]
    assert "300" in settings["Inspection standoff"]
    assert "889.9" in settings["Camera intrinsics"]          # fx from the archived K
    assert "1.26" in settings["Hand-eye calibration"] and "borderline" in settings["Hand-eye calibration"]
    assert "20260629-130945" in settings["Hand-eye calibration"]
    assert settings["Work frame"] == "Tasni Work Frame"
    assert "c48b1ab" in settings["Software revision"]
    assert "1.26" in "\n".join(paragraphs)                    # the floor, in the method text


def test_the_draft_breaks_down_what_one_inspection_costs(tmp_path):
    """Requirement 3 is a cycle time; a reviewer will ask what it is made of."""
    pytest.importorskip("docx")
    from tasni.modules.extrusion.paper_docx import build_paper_docx

    root = tmp_path / "runs" / "extrusion"
    trial_id = MeasureSession.create(root, auto_plan(), note="rings").trial_id
    _take_with_provenance(root, trial_id)

    out = build_paper_docx(root, trial_id, tmp_path / "draft.docx", embed_figures=False)

    tables = _docx_content(out)[1]
    timing = next(t for t in tables if t[0][0] == "Stage")
    stages = {row[0]: row for row in timing[1:]}
    assert "Move to the inspection pose" in stages
    assert "6100" in stages["Move to the inspection pose"][1]
    assert "Whole inspection excursion" in stages
    assert "14460" in stages["Whole inspection excursion"][1]


def test_the_draft_states_its_own_limitations(tmp_path):
    """The claims this experiment cannot support, written down before review."""
    pytest.importorskip("docx")
    from tasni.modules.extrusion.paper_docx import build_paper_docx

    root = tmp_path / "runs" / "extrusion"
    trial_id = MeasureSession.create(root, auto_plan(), note="rings").trial_id
    _take_with_provenance(root, trial_id)

    text = "\n".join(_docx_content(
        build_paper_docx(root, trial_id, tmp_path / "d.docx", embed_figures=False))[0])

    assert "between layers" in text and "during deposition" in text
    assert "Limitations" in text
    assert "reprocess" in text.lower()          # the archive claim, with its evidence


# ------------------------------------------------ paired detection error (2026-08-29)
# The steel rule measures the shift FROM WHERE THE RING SAT, so the chain's error
# must be scored the same way: against the ring's own last measured position
# before it was moved, not against the plan centre (which folds the operator's
# "placed true" bias into what the paper calls the chain's error).


def _stacked_then_shifted(root, *, bias=(1.5, -0.8), shift=(10.2, 0.1), typed=(10, 0)):
    """Ring 3 placed 'true' with a placement bias, measured 3x, then displaced."""
    trial_id = MeasureSession.create(root, auto_plan(), note="rings").trial_id
    for take in (1, 2, 3):
        _write_take(root, trial_id, 3, take, offset=None, measured_offset=list(bias),
                    rms=1.2, mean_abs=1.0, maximum=1.8, acq_ms=1000, phase="stacked true")
    after = [bias[0] + shift[0], bias[1] + shift[1]]
    for take in (4, 5, 6):
        _write_take(root, trial_id, 3, take, offset=list(typed), measured_offset=after,
                    rms=7.1, mean_abs=6.4, maximum=10.2, acq_ms=1000, phase="top ring shifted")
    return trial_id


def test_the_introduced_shift_is_scored_against_the_rings_own_position_before_it_moved(tmp_path):
    root = tmp_path / "runs" / "extrusion"
    trial_id = _stacked_then_shifted(root)

    by_name = {c["condition"]: c for c in paper_summary(root, trial_id)["conditions"]}
    shifted = by_name["layer 3 - top ring shifted - introduced offset (10, 0) mm"]

    # Against the plan centre the 1.5/-0.8 placement bias pollutes the score ...
    assert shifted["detection_error_mm"]["mean"] == pytest.approx(np.hypot(1.7, -0.7), abs=1e-9)
    # ... paired against the ring's last pre-shift measurement it is the chain alone.
    assert shifted["paired_detection_error_mm"]["n"] == 3
    assert shifted["paired_detection_error_mm"]["mean"] == pytest.approx(np.hypot(0.2, 0.1), abs=1e-9)
    assert shifted["paired_shift_norm_mm"]["mean"] == pytest.approx(np.hypot(10.2, 0.1), abs=1e-9)
    assert shifted["paired_reference_takes"] == [3]
    assert by_name["layer 3 - stacked true"]["paired_detection_error_mm"]["n"] == 0


def test_the_pre_shift_reference_is_the_last_zero_offset_take_before_it_never_after(tmp_path):
    """A ring put back after the shift is a later zero-offset take; it is not the reference."""
    root = tmp_path / "runs" / "extrusion"
    trial_id = _stacked_then_shifted(root)
    _write_take(root, trial_id, 3, 7, offset=None, measured_offset=[4.0, 4.0],
                rms=1.2, mean_abs=1.0, maximum=1.8, acq_ms=1000, phase="stacked true")

    shifted = {c["condition"]: c for c in paper_summary(root, trial_id)["conditions"]}[
        "layer 3 - top ring shifted - introduced offset (10, 0) mm"]

    assert shifted["paired_reference_takes"] == [3]
    assert shifted["paired_detection_error_mm"]["mean"] == pytest.approx(np.hypot(0.2, 0.1), abs=1e-9)


def test_a_shift_with_no_prior_measurement_of_that_layer_is_scored_against_the_plan_centre_only(tmp_path):
    root = tmp_path / "runs" / "extrusion"
    trial_id = MeasureSession.create(root, auto_plan(), note="rings").trial_id
    _write_take(root, trial_id, 3, 1, offset=[10, 0], measured_offset=[10.4, 0.3],
                rms=7.0, mean_abs=6.3, maximum=9.9, acq_ms=1000, phase="top ring shifted")

    summary = paper_summary(root, trial_id)
    shifted = summary["conditions"][0]

    assert shifted["paired_detection_error_mm"]["n"] == 0
    assert shifted["paired_reference_takes"] == []
    assert shifted["detection_error_mm"]["mean"] == pytest.approx(0.5, abs=1e-9)
    assert any("plan centre only" in p for p in summary["prose"])


def test_the_markdown_and_prose_state_the_paired_recovery(tmp_path):
    root = tmp_path / "runs" / "extrusion"
    trial_id = _stacked_then_shifted(root)

    summary = paper_summary(root, trial_id)

    assert "paired detection error (mm)" in summary["markdown"]
    recovered = next(p for p in summary["prose"] if "10.0 mm introduced offset" in p)
    assert "10.20" in recovered and "0.22" in recovered
    assert "own last measured position" in recovered and "take 3" in recovered


def test_the_word_draft_carries_the_paired_detection_error(tmp_path):
    pytest.importorskip("docx")
    from tasni.modules.extrusion.paper_docx import build_paper_docx

    root = tmp_path / "runs" / "extrusion"
    trial_id = _stacked_then_shifted(root)

    out = build_paper_docx(root, trial_id, tmp_path / "draft.docx", embed_figures=False)

    paragraphs, tables = _docx_content(out)
    conditions = next(t for t in tables if t[0][0] == "Condition")
    column = conditions[0].index("Paired detection error (mm)")
    rows = {row[0]: row for row in conditions[1:]}
    assert "0.22" in rows["layer 3 - top ring shifted - introduced offset (10, 0) mm"][column]
    assert rows["layer 3 - stacked true"][column] == "-"
    takes = next(t for t in tables if t[0][0] == "Layer")
    assert "Paired error (mm)" in takes[0]


def test_the_draft_asks_for_a_pre_shift_measurement_when_a_shift_has_none(tmp_path):
    pytest.importorskip("docx")
    from tasni.modules.extrusion.paper_docx import build_paper_docx

    root = tmp_path / "runs" / "extrusion"
    trial_id = MeasureSession.create(root, auto_plan(), note="rings").trial_id
    for take in (1, 2, 3):
        _write_take(root, trial_id, 3, take, offset=[10, 0], measured_offset=[10.4, 0.3],
                    rms=7.0, mean_abs=6.3, maximum=9.9, acq_ms=1000, phase="top ring shifted")

    text = "\n".join(_docx_content(
        build_paper_docx(root, trial_id, tmp_path / "d.docx", embed_figures=False))[0])

    assert "measure the ring in place before displacing it" in text.lower()


def test_the_timing_requirement_counts_live_measurements_not_takes(tmp_path):
    """An offline reprocess has no live acquisition-to-path time, so it owes one more."""
    pytest.importorskip("docx")
    from tasni.modules.extrusion.paper_docx import build_paper_docx

    root = tmp_path / "runs" / "extrusion"
    trial_id = MeasureSession.create(root, auto_plan(), note="rings").trial_id
    _write_take(root, trial_id, 1, 1, offset=None, offset_norm=0.4, rms=0.5, mean_abs=0.4,
                maximum=1.1, acq_ms=1000, phase="noise floor", offline=True)

    text = "\n".join(_docx_content(
        build_paper_docx(root, trial_id, tmp_path / "d.docx", embed_figures=False))[0])

    assert "0 live measurement(s) recorded" in text and "12 more" in text


def test_the_axis_check_take_is_never_the_pre_shift_reference(tmp_path):
    """The 'which way is +X' take moved the ring an untyped amount: not where it sat."""
    root = tmp_path / "runs" / "extrusion"
    trial_id = MeasureSession.create(root, auto_plan(), note="rings").trial_id
    for take in (1, 2, 3):
        _write_take(root, trial_id, 3, take, offset=None, measured_offset=[1.5, -0.8],
                    rms=1.2, mean_abs=1.0, maximum=1.8, acq_ms=1000, phase="stacked true")
    # The throwaway: ring slid ~9 mm to learn the axis, nothing typed.
    _write_take(root, trial_id, 3, 4, offset=None, measured_offset=[1.5 + 9.0, -0.8],
                rms=6.0, mean_abs=5.5, maximum=9.0, acq_ms=1000, phase="axis check")
    for take in (5, 6, 7):
        _write_take(root, trial_id, 3, take, offset=[10, 0], measured_offset=[11.7, -0.7],
                    rms=7.1, mean_abs=6.4, maximum=10.2, acq_ms=1000, phase="top ring shifted")

    shifted = {c["condition"]: c for c in paper_summary(root, trial_id)["conditions"]}[
        "layer 3 - top ring shifted - introduced offset (10, 0) mm"]

    assert shifted["paired_reference_takes"] == [3]
    assert shifted["paired_detection_error_mm"]["mean"] == pytest.approx(np.hypot(0.2, 0.1), abs=1e-9)


# ------------- the depth must be of the pose it was taken at (cell, 2026-08-29)

def test_depth_plane_check_scales_raw_depth_words_by_the_frames_own_unit():
    """Task 9 review, Critical 1: ``depth`` is raw camera WORDS, not millimetres.

    Protocol 2 native depth is 0.1 mm/word. Left at the unit-blind default this
    reads a real ~312 mm standoff as ~3120 mm, fails loudly, and blames the work
    frame or a frozen camera -- the wrong cause -- on every single measure and
    characterize take.
    """
    from tasni.modules.extrusion.measure import depth_plane_check
    # Rotation matters here: the camera's +Z runs OUT of the lens, so an identity
    # rotation faces away from the plane. diag(1,-1,-1) is the straight-down pose.
    T = np.eye(4)
    T[:3, :3] = np.diag([1.0, -1.0, -1.0])
    T[2, 3] = 312.0                                   # camera 312 mm above the plane
    depth_words = np.full((8, 8), 3120, np.uint16)     # 3120 words * 0.1 mm/word = 312 mm

    scaled = depth_plane_check(depth_words, T, ExtrusionConfig(), unit_mm=0.1)
    assert scaled["observed_depth_mm"] == pytest.approx(312.0)
    assert scaled["agrees"] is True

    # The old, unit-blind call (no unit_mm) reads the same words as 3120 mm --
    # this is the regression the review measured on the checkout.
    unscaled = depth_plane_check(depth_words, T, ExtrusionConfig())
    assert unscaled["observed_depth_mm"] == pytest.approx(3120.0)
    assert unscaled["agrees"] is False


class StaleThenFreshCamera:
    """The Jetson's temporal depth filter blends across frames.

    After the arm moves, the first depth it emits is still weighted toward the
    PREVIOUS distance. On the cell this returned the parked height (447 mm) with
    a correct, freshly captured colour of the board at 312 mm -- geometry that
    could not both be true, and which put every point 135 mm outside the search
    region.
    """

    def __init__(self, stale_mm=640, fresh_mm=500, stale_frames=1):
        self.stale_mm, self.fresh_mm = stale_mm, fresh_mm
        self.stale_frames = stale_frames
        self.grabs = 0

    def grab(self, **kwargs):
        self.grabs += 1
        # The readiness grab happens before the move and is not counted against
        # the stale run: it is the capture after the move that must be fresh.
        value = self.stale_mm if self.grabs <= self.stale_frames + 1 else self.fresh_mm
        from tasni.core.camera import Frame
        return Frame(color=np.zeros((16, 16, 3), np.uint8),
                     depth=np.full((16, 16), value, np.uint16), timestamp=1.0,
                     geometry=gf.aligned(syn.K_720P, (16, 16)))

    def stream(self, **kwargs):
        from contextlib import nullcontext
        from types import SimpleNamespace
        camera = self
        return nullcontext(SimpleNamespace(
            read=lambda **read_kwargs: camera.grab(**read_kwargs)))


def test_a_depth_frame_from_the_wrong_pose_is_grabbed_again(tmp_path, monkeypatch):
    """A stale depth is silently wrong, so it must not be measured."""
    svc, rdk, _ = measure_env(tmp_path, monkeypatch)
    svc.camera = StaleThenFreshCamera()
    plan = auto_plan()
    session = MeasureSession.create(tmp_path / "runs" / "extrusion", plan)

    out = RingMeasureJob(svc, plan, session, 1, annotation={}, check_collisions=True)(Ctx())

    manifest = json.loads((Path(out["layer_dir"]) / "manifest.json").read_text())
    check = manifest["processing"]["depth_plane_check"]
    assert check["retries"] == 1, "the stale frame should have been thrown away"
    assert check["camera_z_mm"] == pytest.approx(506.0)
    assert check["observed_depth_mm"] == pytest.approx(500.0)


def test_a_depth_frame_that_never_matches_the_pose_fails_loudly(tmp_path, monkeypatch):
    """Better a refused measurement than a plausible wrong one."""
    svc, rdk, _ = measure_env(tmp_path, monkeypatch)
    svc.camera = StaleThenFreshCamera(stale_frames=99)
    plan = auto_plan()
    session = MeasureSession.create(tmp_path / "runs" / "extrusion", plan)

    with pytest.raises(RuntimeError) as failure:
        RingMeasureJob(svc, plan, session, 1, annotation={}, check_collisions=True)(Ctx())

    message = str(failure.value)
    assert "640" in message and "506" in message, message
    # The frames were byte-identical, which is what a stalled stream looks like:
    # the operator is told that, and the one command that clears it.
    assert "frozen" in message.lower(), message
    assert "jetson_deploy.py restart" in message, message
    assert rdk.events[-1] == ("move-joints", START_JOINTS)          # still comes home


# -- one press, several takes ------------------------------------------------
# Two questions the run asks separately, because they cost differently: how
# repeatably the CHAIN sees a ring that has not moved (frames with the arm
# parked, seconds each), and how repeatably the ARM comes back to look at it
# (whole trips out and back, a trip each). Before this, both were three presses
# of the same button and the difference was not recorded anywhere.

def test_parked_repeats_take_several_frames_on_one_trip(tmp_path, monkeypatch):
    """repeats=3 leaves the path ONCE and banks three takes from that pose."""
    svc, rdk, camera = measure_env(tmp_path, monkeypatch)
    plan = auto_plan()
    session = MeasureSession.create(tmp_path / "runs" / "extrusion", plan)

    out = RingMeasureJob(svc, plan, session, 1, check_collisions=False, repeats=3)(Ctx())

    starts = [e for e in rdk.events if e[0] == "start"]
    homes = [e for e in rdk.events if e == ("move-joints", START_JOINTS)]
    assert len(starts) == 1, "a parked repeat must not re-run the inspection move"
    assert len(homes) == 1, "the arm comes home once, after the last frame"
    assert camera.grabs == 4                       # readiness + three measurements
    assert session.takes == {1: 3}
    assert out["takes_recorded"] == [1, 2, 3]
    names = sorted(p.parent.name for p in (tmp_path / "runs" / "extrusion").glob(
        "*/layer-*/manifest.json"))
    assert names == ["layer-001", "layer-001-take02", "layer-001-take03"]


def test_a_shared_trip_is_never_priced_as_one_take_s_excursion(tmp_path, monkeypatch):
    """The cycle figure is the cost of leaving the path for ONE measurement.

    Three frames from one trip cost one trip between them. Letting each claim an
    inspection_cycle_ms would divide the excursion's real price by three in the
    one statistic the paper quotes as scan-to-feedback turnaround.
    """
    svc, rdk, camera = measure_env(tmp_path, monkeypatch)
    plan = auto_plan()
    session = MeasureSession.create(tmp_path / "runs" / "extrusion", plan)

    RingMeasureJob(svc, plan, session, 1, check_collisions=False, repeats=3)(Ctx())

    root = tmp_path / "runs" / "extrusion" / session.trial_id
    timings = [json.loads((root / name / "manifest.json").read_text())["processing"]["timings_ms"]
               for name in ("layer-001", "layer-001-take02", "layer-001-take03")]
    assert not any("inspection_cycle_ms" in t for t in timings)
    # The move belongs to the frame that followed it, the return to the last.
    assert "move_to_pose_ms" in timings[0] and "settle_ms" in timings[0]
    assert not any("move_to_pose_ms" in t for t in timings[1:])
    assert "return_ms" in timings[2] and "return_ms" not in timings[0]
    # ...and every frame still carries what it really measured.
    assert all(t["capture_ms"] >= 0 and t["acquisition_to_path_ms"] > 0 for t in timings)


def test_excursions_repeat_the_whole_trip_unattended(tmp_path, monkeypatch):
    """excursions=5 is five presses of the noise-floor button, without pressing.

    The arm must genuinely leave and come back each time -- that re-approach is
    the thing being measured -- and each take is priced as a whole excursion.
    """
    svc, rdk, camera = measure_env(tmp_path, monkeypatch)
    plan = auto_plan()
    session = MeasureSession.create(tmp_path / "runs" / "extrusion", plan)

    out = RingMeasureJob(svc, plan, session, 1, annotation={"phase": "noise floor"},
                         check_collisions=False, excursions=5)(Ctx())

    assert len([e for e in rdk.events if e[0] == "start"]) == 5
    assert len([e for e in rdk.events if e == ("move-joints", START_JOINTS)]) == 5
    assert out["takes_recorded"] == [1, 2, 3, 4, 5]
    root = tmp_path / "runs" / "extrusion" / session.trial_id
    cycles = [json.loads(p.read_text())["processing"]["timings_ms"].get("inspection_cycle_ms")
              for p in sorted(root.glob("layer-001*/manifest.json"))]
    assert len(cycles) == 5 and all(c is not None for c in cycles)


def test_unattended_excursions_stop_after_an_archived_invalid_take(tmp_path, monkeypatch):
    """A failed gate must be visible and home before another robot trip starts."""
    svc, rdk, camera = measure_env(tmp_path, monkeypatch, side_photo=True)
    plan = auto_plan()
    session = MeasureSession.create(tmp_path / "runs" / "extrusion", plan)

    def invalid(**kwargs):
        result = fake_measure_processing(**kwargs)
        result.metrics = result.metrics.model_copy(update={
            "valid": False, "warnings": ["maximum angular gap 200 deg"]})
        return result

    monkeypatch.setattr(processing_mod, "process_observation", invalid)
    ctx = Ctx()
    out = RingMeasureJob(svc, plan, session, 1, annotation={"phase": "noise floor"},
                         check_collisions=False, excursions=5)(ctx)

    assert out["stopped_early"] is True
    assert out["invalid_batch"] is True
    assert out["side_view"] is None
    assert camera.witness_grabs == 0
    assert out["takes_recorded"] == [1]
    assert len([e for e in rdk.events if e[0] == "start"]) == 1
    assert rdk.events[-1] == ("move-joints", START_JOINTS)
    assert ctx.checkpoints == [("extrusion_take", {
        "trial_id": session.trial_id, "layer_index": 1, "take": 1,
        "valid": False, "layer_name": "layer-001"})]
    assert any("stopped before 4 remaining" in line for line in ctx.logs)


def test_each_take_records_which_trip_it_came_from(tmp_path, monkeypatch):
    """Three frames of one trip and three re-approaches must be tellable apart.

    The difference between those two spreads IS the finding the noise floor
    supports, so a reader of the archive alone has to be able to separate them.
    """
    svc, rdk, camera = measure_env(tmp_path, monkeypatch)
    plan = auto_plan()
    session = MeasureSession.create(tmp_path / "runs" / "extrusion", plan)

    RingMeasureJob(svc, plan, session, 1, check_collisions=False,
                   excursions=2, repeats=2)(Ctx())

    root = tmp_path / "runs" / "extrusion" / session.trial_id
    stamps = [json.loads(p.read_text())["provenance"]
              for p in sorted(root.glob("layer-001*/manifest.json"))]
    assert [(s["excursion_index"], s["repeat_index"]) for s in stamps] == [
        (1, 1), (1, 2), (2, 1), (2, 2)]
    assert all(s["repeats_in_excursion"] == 2 for s in stamps)


def test_a_batch_keeps_the_takes_it_already_measured(tmp_path, monkeypatch):
    """A failure on trip three does not throw away trips one and two.

    Unattended batches are the point of excursions; losing a whole run to the
    last frame would make them worse than pressing the button five times.
    """
    svc, rdk, camera = measure_env(tmp_path, monkeypatch)
    plan = auto_plan()
    session = MeasureSession.create(tmp_path / "runs" / "extrusion", plan)
    calls = {"n": 0}

    def fail_on_the_third(**kwargs):
        calls["n"] += 1
        if calls["n"] == 3:
            raise RuntimeError("branch guard exhausted")
        return fake_measure_processing(**kwargs)

    monkeypatch.setattr(processing_mod, "process_observation", fail_on_the_third)
    with pytest.raises(RuntimeError, match="measurement invalid"):
        RingMeasureJob(svc, plan, session, 1, check_collisions=False, excursions=5)(Ctx())

    assert session.takes == {1: 3}, "the failed take is still numbered and archived"
    valid = [r for r in session.records if r["valid"]]
    assert len(valid) == 2, "the two good takes survive for the next press to continue from"
    # And the arm is home, not parked over the ring.
    assert rdk.events[-1] == ("move-joints", START_JOINTS)


def test_the_api_carries_the_batch_counts_and_caps_them(tmp_path, monkeypatch):
    """A mistyped count must not commit the cell to unattended motion for a quarter hour."""
    from pydantic import ValidationError

    from tasni.modules.extrusion.module import MeasureLayerBody

    body = MeasureLayerBody(fingerprint="f", layer_index=1)
    assert (body.repeats, body.excursions) == (1, 1)
    assert MeasureLayerBody(fingerprint="f", layer_index=1, repeats=3).repeats == 3
    for bad in ({"repeats": 0}, {"excursions": 11}, {"repeats": -1}):
        with pytest.raises(ValidationError):
            MeasureLayerBody(fingerprint="f", layer_index=1, **bad)


def test_the_summary_separates_the_camera_s_repeatability_from_the_robot_s(tmp_path, monkeypatch):
    """One "repeatability" over both would credit the robot's scatter to the camera.

    Parked frames scatter by the sensing chain alone; takes that each cost a
    trip out and back scatter by the chain AND the re-approach. The paper wants
    to say what re-approaching costs, so the two must be pooled separately.
    """
    from tasni.modules.extrusion.measure import capture_style, centre_spread, paper_summary

    svc, rdk, camera = measure_env(tmp_path, monkeypatch)
    plan = auto_plan()
    root = tmp_path / "runs" / "extrusion"
    session = MeasureSession.create(root, plan)
    # A wandering fitted centre, so the two pools cannot come out identical by
    # accident: the noise-floor trips scatter ten times as far as the parked ones.
    centres = iter([(200.0, 150.0), (210.0, 150.0), (190.0, 150.0),     # 3 trips
                    (200.0, 150.0), (201.0, 150.0), (199.0, 150.0)])    # 3 parked frames

    def wander(**kwargs):
        out = fake_measure_processing(**kwargs)
        out.metrics.measured_center_mm = next(centres)
        return out

    monkeypatch.setattr(processing_mod, "process_observation", wander)
    RingMeasureJob(svc, plan, session, 1, annotation={"phase": "noise floor"},
                   check_collisions=False, excursions=3)(Ctx())
    RingMeasureJob(svc, plan, session, 1, annotation={"phase": "re-placed"},
                   check_collisions=False, repeats=3)(Ctx())

    summary = paper_summary(root, session.trial_id)
    by_name = {c["condition"]: c for c in summary["conditions"]}
    trips = by_name["layer 1 - noise floor"]
    parked = by_name["layer 1 - re-placed"]
    assert trips["capture"] == "re-approach" and parked["capture"] == "parked"
    assert centre_spread([]) == {"n": 0, "rms_mm": None, "max_mm": None}
    assert capture_style([{}]) == "single"

    repeat = summary["repeatability_mm"]
    assert repeat["sensing"]["takes"] == 3 and repeat["re_approach"]["takes"] == 3
    assert repeat["sensing"]["rms_mm"] == pytest.approx(1.0, abs=0.01)
    assert repeat["re_approach"]["rms_mm"] == pytest.approx(10.0, abs=0.01)
    # And it is stated, not merely computed: the paper is written from the prose.
    said = " ".join(summary["prose"])
    assert "Sensing repeatability" in said and "Re-approach repeatability" in said


def test_old_takes_without_a_trip_stamp_read_as_their_own_excursion(tmp_path, monkeypatch):
    """Every take archived before the batch split WAS one press, so one trip.

    Reading them as parked would retro-credit the chain with a repeatability it
    was never measured to have.
    """
    from tasni.modules.extrusion.measure import capture_style

    legacy = [{"metrics": {"valid": True, "measured_center_mm": [200.0, 150.0]}},
              {"metrics": {"valid": True, "measured_center_mm": [201.0, 150.0]}}]
    assert capture_style(legacy) == "re-approach"


def test_a_ring_placed_displaced_is_paired_against_the_ring_it_sits_on(tmp_path, monkeypatch):
    """Ring 4 arrives already off-centre, so it has no earlier self to pair with.

    The rule measured how far it sits from ring 3, so that is what it must be
    scored against. Scored against the plan centre instead, the whole stack's
    hand-placement error is charged to the sensing chain.
    """
    from tasni.modules.extrusion.measure import paper_summary

    svc, rdk, camera = measure_env(tmp_path, monkeypatch)
    plan = auto_plan(layers=4)
    root = tmp_path / "runs" / "extrusion"
    session = MeasureSession.create(root, plan)
    # The stack was hand-placed 3 mm off the plan centre; ring 4 then sits a
    # further 10 mm out. The chain must report 10, not 13.
    centres = iter([(203.0, 150.0)] * 3 + [(213.0, 150.0)] * 3)

    def at_position(**kwargs):
        out = fake_measure_processing(**kwargs)
        centre = next(centres)
        out.metrics.measured_center_mm = centre
        out.metrics.center_offset_mm = (centre[0] - 200.0, centre[1] - 150.0)
        return out

    monkeypatch.setattr(processing_mod, "process_observation", at_position)
    RingMeasureJob(svc, plan, session, 3, annotation={"phase": "stacked true"},
                   check_collisions=False, repeats=3)(Ctx())
    RingMeasureJob(svc, plan, session, 4,
                   annotation={"phase": "top ring shifted",
                               "introduced_offset_mm": [10, 0]},
                   check_collisions=False, repeats=3)(Ctx())

    summary = paper_summary(root, session.trial_id)
    shifted = next(c for c in summary["conditions"] if c["introduced_norm_mm"] > 0)
    assert shifted["paired_reference_layers"] == [3], "paired against the ring beneath it"
    assert shifted["paired_against_layer_below"] is True
    assert shifted["paired_shift_norm_mm"]["mean"] == pytest.approx(10.0, abs=0.01)
    assert shifted["paired_detection_error_mm"]["mean"] == pytest.approx(0.0, abs=0.01)
    # The unpaired score carries the stack's own 3 mm placement error.
    assert shifted["detection_error_mm"]["mean"] == pytest.approx(3.0, abs=0.01)
    assert any("the ring it was stacked on" in p for p in summary["prose"])


def test_a_ring_that_was_slid_still_pairs_against_its_own_earlier_take(tmp_path, monkeypatch):
    """The same-layer reference stays first choice; the layer below is a fallback.

    A ring measured true and THEN displaced has its own undisplaced position on
    record, which is a tighter reference than the ring beneath it.
    """
    from tasni.modules.extrusion.measure import pre_shift_reference

    same_layer = {"layer_index": 3, "take": 2, "annotation": {},
                  "metrics": {"valid": True, "measured_center_mm": [201.0, 150.0]}}
    below = {"layer_index": 2, "take": 9, "annotation": {},
             "metrics": {"valid": True, "measured_center_mm": [200.0, 150.0]}}
    shifted = {"layer_index": 3, "take": 5,
               "annotation": {"introduced_offset_mm": [10, 0]},
               "metrics": {"valid": True, "measured_center_mm": [211.0, 150.0]}}

    assert pre_shift_reference(shifted, [below, same_layer, shifted]) is same_layer
    # With no take of its own, it falls back to the ring beneath.
    assert pre_shift_reference(shifted, [below, shifted]) is below
    # Layer 1 has nothing beneath it, so it stays unpaired rather than inventing one.
    lonely = {**shifted, "layer_index": 1}
    assert pre_shift_reference(lonely, [below, lonely]) is None

def test_the_side_photo_goes_out_and_back_through_the_approach_target(tmp_path, monkeypatch):
    """Neutral -> approach -> side, and side -> approach -> neutral.

    The approach target is not a nicety: the direct joint move between the
    neutral pose and the side pose is the one that sweeps the arm through the
    things standing around the cell. Skipping it on the way BACK would bump into
    them just as hard as skipping it on the way out.
    """
    svc, rdk, camera = measure_env(tmp_path, monkeypatch, side_photo=True)
    plan = auto_plan()
    session = MeasureSession.create(tmp_path / "runs" / "extrusion", plan)

    out = RingMeasureJob(svc, plan, session, 1, check_collisions=False, repeats=2)(Ctx())

    route = [e for e in rdk.events if e[0] in {"move-target", "move-joints"}]
    # By STORED JOINTS, not by target item: a cartesian MoveJ resolves against
    # whatever tool/frame the last take left active, which on the cell sent the
    # arm somewhere else entirely (137.8 s excursion, 2.7 s is an inspection move).
    assert route[-4:] == [("move-joints", SIDE_APPROACH_JOINTS),
                          ("move-joints", SIDE_JOINTS),
                          ("move-joints", SIDE_APPROACH_JOINTS),
                          ("move-joints", START_JOINTS)]
    assert out["side_view"]["captured"] is True
    assert camera.witness_grabs == 1, "the side photo is RGB only -- it measures nothing"


def test_a_cartesian_only_side_target_still_moves_but_says_it_is_ambiguous(tmp_path,
                                                                          monkeypatch):
    """No stored joints = no unambiguous move. Do it, but do not do it silently."""
    svc, rdk, camera = measure_env(tmp_path, monkeypatch, side_photo=True)
    rdk.taught_joints["SideCapture"] = None
    plan = auto_plan()
    session = MeasureSession.create(tmp_path / "runs" / "extrusion", plan)
    ctx = Ctx()

    out = RingMeasureJob(svc, plan, session, 1, check_collisions=False)(ctx)

    assert out["side_view"]["captured"] is True
    assert ("move-target", "SideCapture") in rdk.events
    assert any("stores no joints" in m and "SideCapture" in m for m in ctx.logs), ctx.logs


def test_the_side_photo_is_taken_once_per_press_and_lands_on_the_last_take(tmp_path, monkeypatch):
    """The ring does not move between the frames of one capture, so one photo.

    It is archived beside the take it belongs to, so a reader of the archive
    finds the picture next to the numbers rather than in a separate pile.
    """
    svc, rdk, camera = measure_env(tmp_path, monkeypatch, side_photo=True)
    plan = auto_plan()
    root = tmp_path / "runs" / "extrusion"
    session = MeasureSession.create(root, plan)

    RingMeasureJob(svc, plan, session, 1, check_collisions=False, repeats=3)(Ctx())

    assert len([e for e in rdk.events if e == ("move-joints", SIDE_JOINTS)]) == 1
    trial = root / session.trial_id
    assert (trial / "layer-001-take03" / "side.png").is_file()
    assert not (trial / "layer-001" / "side.png").exists()
    manifest = json.loads((trial / "layer-001-take03" / "manifest.json").read_text())
    assert manifest["side_view"]["captured"] is True
    assert manifest["side_view"]["image_file"] == "side.png"
    assert manifest["side_view"]["target"] == "SideCapture"
    assert manifest["side_view"]["approach_target"] == "TowardsSideCapture"
    # Not folded into what the paper says an inspection costs.
    assert "side" not in json.dumps(manifest["processing"]["timings_ms"])
    assert session.records[-1]["side_view"]["captured"] is True


def test_a_missing_taught_target_skips_the_photo_and_never_moves(tmp_path, monkeypatch):
    """A station without the targets must still be able to measure.

    And it must not set off toward a target that is not there: the skip has to
    happen BEFORE any motion, not as a failed move.
    """
    svc, rdk, camera = measure_env(tmp_path, monkeypatch, side_photo=True)
    rdk.absent.add("TowardsSideCapture")
    plan = auto_plan()
    session = MeasureSession.create(tmp_path / "runs" / "extrusion", plan)

    out = RingMeasureJob(svc, plan, session, 1, check_collisions=False)(Ctx())

    assert out["valid"] is True, "the measurement stands"
    assert not any(e[0] == "move-target" for e in rdk.events)
    side = out["side_view"]
    assert side["captured"] is False
    assert "TowardsSideCapture" in side["error"] and "not in the station" in side["error"]
    # Archived anyway: the manifest has to be able to say why there is no photo.
    manifest = json.loads((Path(out["layer_dir"]) / "manifest.json").read_text())
    assert manifest["side_view"]["captured"] is False and manifest["side_view"]["error"]


def test_a_failure_at_the_side_pose_still_retraces_and_keeps_the_measurement(tmp_path, monkeypatch):
    """An arm left parked out at the side pose is the worst outcome available.

    So the return leg runs even when the capture failed -- and the measurement,
    which is already on disk, is not put at risk for a figure.
    """
    svc, rdk, camera = measure_env(tmp_path, monkeypatch, side_photo=True)
    plan = auto_plan()
    session = MeasureSession.create(tmp_path / "runs" / "extrusion", plan)

    def dead_camera(**kwargs):
        if kwargs.get("color_only"):
            raise RuntimeError("camera stopped answering")
        return FakeCamera.grab(camera, **kwargs)

    monkeypatch.setattr(camera, "grab", dead_camera)
    out = RingMeasureJob(svc, plan, session, 1, check_collisions=False)(Ctx())

    assert out["valid"] is True
    assert out["side_view"]["captured"] is False
    assert "camera stopped answering" in out["side_view"]["error"]
    route = [e for e in rdk.events if e[0] in {"move-target", "move-joints"}]
    assert route[-2:] == [("move-joints", SIDE_APPROACH_JOINTS),
                          ("move-joints", START_JOINTS)]


def test_the_side_photo_is_part_of_the_protocol_by_default(tmp_path, monkeypatch):
    """It is a step of the run, not an opt-in: the paper needs one per layer."""
    from tasni.core.config import AppConfig
    from tasni.modules.extrusion.module import MeasureLayerBody

    config = AppConfig().extrusion
    assert config.measure_depth_fusion_frames == 5
    assert config.side_capture_enabled is True
    assert config.side_capture_target == "SideCapture"
    assert config.side_capture_approach_target == "TowardsSideCapture"
    assert MeasureLayerBody(fingerprint="f", layer_index=1).side_photo is None


def test_the_side_photo_is_served_to_the_browser_and_listed(tmp_path, monkeypatch):
    """A photo the operator cannot see is a photo they cannot check was framed."""
    svc, rdk, camera = measure_env(tmp_path, monkeypatch, side_photo=True)
    plan = auto_plan()
    root = tmp_path / "runs" / "extrusion"
    session = MeasureSession.create(root, plan)
    RingMeasureJob(svc, plan, session, 1, check_collisions=False)(Ctx())
    monkeypatch.setattr(extrusion_module, "REPO_ROOT", tmp_path)
    client = TestClient(create_app(AppConfig()))

    served = client.get(f"/api/modules/extrusion/trials/{session.trial_id}"
                        "/layers/layer-001/files/side.png")
    assert served.status_code == 200 and served.headers["content-type"] == "image/png"

    listed = client.get("/api/modules/extrusion/trials").json()
    take = next(t for i in listed["trials"] if i["trial_id"] == session.trial_id
                for t in i["layers"])
    assert take["has_side_view"] is True
    assert take["side_view"]["captured"] is True

    # A path that is not on the whitelist is still refused.
    assert client.get(f"/api/modules/extrusion/trials/{session.trial_id}"
                      "/layers/layer-001/files/depth.npy").status_code == 404


# ------------------------------------------------- substrate health gating
# `separation_margin_mm` -- spec section 4's one derived "is segmentation
# healthy here" number -- was computed, written into the report, and read by
# nothing. Measured 2026-08-31 on the cell: a take whose deposit median sat
# BELOW the board's own p99 (margin -0.119 mm, sigma 0.866, floor pinned on its
# 2.0 mm clamp) returned `valid: true` with `warnings: []`, and reported a bead
# 41% fatter than the same ring measured an hour earlier. Every honest take in
# the 2026-08-30 archive sits at margin 2.05-2.18 with sigma 0.515-0.568 and a
# floor of 1.55-1.70, so there is ~2 mm of headroom on both sides of the fault
# line; this is a gate, not a tuning parameter.

from tasni.modules.extrusion.processing import substrate_health  # noqa: E402


def _substrate(**over):
    """One golden layer-1 take's substrate block (2026-08-30 archive, layer-001)."""
    block = {"source": "fitted_plane", "sigma_mm": 0.5635, "floor_mm": 1.691,
             "substrate_p99_mm": 1.364, "separation_margin_mm": 2.182}
    block.update(over)
    return block


def test_substrate_health_passes_a_take_from_the_golden_archive():
    fault, warnings = substrate_health(_substrate(), ExtrusionConfig())
    assert fault is None
    assert warnings == []


def test_substrate_health_faults_when_the_deposit_sits_below_the_board():
    """The 2026-08-31 17:12 cell take. A negative margin means the deposit's
    median height is under the board's 99th percentile -- the two populations
    are not separable, so whatever came out is not a measurement."""
    fault, _ = substrate_health(
        _substrate(separation_margin_mm=-0.119, sigma_mm=0.8657, floor_mm=2.0,
                   substrate_p99_mm=3.535),
        ExtrusionConfig())
    assert fault is not None
    assert "-0.119" in fault                    # names the number it refused on
    assert "separation" in fault.lower()


def test_substrate_health_warns_on_a_thin_margin_without_faulting():
    """Above the fault line but below the golden band: report it, measure it."""
    fault, warnings = substrate_health(
        _substrate(separation_margin_mm=1.0), ExtrusionConfig())
    assert fault is None
    assert any("separation" in w.lower() for w in warnings), warnings


def test_substrate_health_warns_when_the_derived_floor_is_pinned_on_its_clamp():
    """A pinned floor means 3*sigma exceeded the clamp ceiling -- the surface is
    noisier than the floor can express, so the floor is no longer derived from
    this frame at all. Structural, and it fires on nothing in the archive."""
    config = ExtrusionConfig()
    ceiling = float(tuple(config.substrate_floor_clamp_mm)[1])
    fault, warnings = substrate_health(_substrate(floor_mm=ceiling), config)
    assert fault is None
    assert any("clamp" in w.lower() for w in warnings), warnings


def test_substrate_health_warns_on_a_substrate_noisier_than_the_archive():
    fault, warnings = substrate_health(_substrate(sigma_mm=0.8657), ExtrusionConfig())
    assert fault is None
    assert any("sigma" in w.lower() for w in warnings), warnings


def test_substrate_health_cannot_fault_on_a_margin_it_does_not_have():
    """`separation_margin_mm` is None when the substrate had no p99 to subtract.
    Absent evidence is not evidence of a fault."""
    fault, warnings = substrate_health(
        _substrate(separation_margin_mm=None, substrate_p99_mm=None),
        ExtrusionConfig())
    assert fault is None
    assert not any("separation" in w.lower() for w in warnings), warnings


def test_the_health_gate_is_wired_into_the_measurement_path():
    """The pure function above is worthless if nothing calls it. A clean
    synthetic ring measures fine on the shipped config and must REFUSE when the
    fault line is raised above its own margin -- which can only happen if
    process_observation consults the gate."""
    pytest.importorskip("open3d")
    T = syn.inspection_camera_T([CENTER[0], CENTER[1], 6.0], 300.0)
    ring = syn.RingSpec(40.0, 10.0, CENTER, height_fn=syn.flat(6.0))
    scene = np.vstack((syn.plane_points(center_xy_mm=CENTER), ring.surface_points()))
    depth = syn.render_depth(scene, T, noise_mm=0.3)
    plan = scene_plan(radius=40.0, bead=10.0, layer_height=6.0, center=CENTER)

    def run(config):
        return process_observation(depth=depth, geometry=syn.geometry(),
                                   T_work_camera=T, plan=plan, layer=plan.layers[0],
                                   config=config)

    healthy = run(ExtrusionConfig())
    assert healthy.report["valid"]
    margin = healthy.report["substrate"]["separation_margin_mm"]
    assert margin > 0.0

    with pytest.raises(RuntimeError, match="separation"):
        run(ExtrusionConfig(substrate_min_separation_mm=margin + 1.0))
